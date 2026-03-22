import streamlit as st
import sqlite3
import hashlib
import os
import numpy as np
from PIL import Image, ImageFilter, ImageEnhance, ImageOps
import io
import shutil
import zipfile
import base64
from datetime import date, timedelta, datetime
import calendar
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# CONFIGURATION
# ============================================================
st.set_page_config(
    page_title="Cabinet Traumatologie & Orthopédie",
    page_icon="🦴", layout="wide",
    initial_sidebar_state="expanded"
)

BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
DB_PATH    = os.path.join(BASE_DIR, "data", "cabinet.db")
RADIO_DIR  = os.path.join(BASE_DIR, "data", "radios")
EXPORT_DIR = os.path.join(BASE_DIR, "data", "exports")

for d in [os.path.join(BASE_DIR,"data"), RADIO_DIR, EXPORT_DIR]:
    os.makedirs(d, exist_ok=True)

# ============================================================
# CSS GLOBAL
# ============================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
* { font-family: 'Inter', sans-serif; }
.main-header { background:linear-gradient(135deg,#1a3a5c 0%,#2d6a9f 100%); padding:1.5rem 2rem;
    border-radius:12px; margin-bottom:1.5rem; color:white; display:flex; align-items:center; gap:1rem; }
.main-header h1 { margin:0; font-size:1.6rem; font-weight:700; }
.main-header p  { margin:0; font-size:0.85rem; opacity:0.85; }
.role-badge { background:rgba(255,255,255,0.2); border:1px solid rgba(255,255,255,0.3);
    padding:0.3rem 0.8rem; border-radius:20px; font-size:0.75rem; font-weight:600; margin-left:auto; }
.card { background:white; border:1px solid #e5e7eb; border-radius:12px;
    padding:1.5rem; margin-bottom:1rem; box-shadow:0 1px 3px rgba(0,0,0,0.06); }
.stat-card { background:linear-gradient(135deg,#1a3a5c,#2d6a9f); color:white;
    border-radius:12px; padding:1.2rem; text-align:center; }
.stat-card .number { font-size:2rem; font-weight:700; }
.stat-card .label  { font-size:0.8rem; opacity:0.85; margin-top:0.2rem; }
.stat-card.green  { background:linear-gradient(135deg,#065f46,#059669); }
.stat-card.orange { background:linear-gradient(135deg,#92400e,#d97706); }
.stat-card.purple { background:linear-gradient(135deg,#4c1d95,#7c3aed); }
.badge { display:inline-block; padding:0.2rem 0.6rem; border-radius:20px; font-size:0.72rem; font-weight:600; }
.badge-blue   { background:#dbeafe; color:#1d4ed8; }
.badge-green  { background:#d1fae5; color:#065f46; }
.badge-red    { background:#fee2e2; color:#991b1b; }
.badge-orange { background:#ffedd5; color:#9a3412; }
.badge-gray   { background:#f3f4f6; color:#374151; }
.rdv-card { border-left:4px solid #2d6a9f; background:#f8fafc;
    padding:0.8rem 1rem; border-radius:0 8px 8px 0; margin-bottom:0.5rem; }
.section-title { font-size:1.1rem; font-weight:600; color:#1a3a5c;
    border-bottom:2px solid #dbeafe; padding-bottom:0.5rem; margin-bottom:1rem; }
.alert-box { padding:0.8rem 1rem; border-radius:8px; margin-bottom:0.8rem; font-size:0.9rem; }
.alert-info    { background:#dbeafe; border-left:4px solid #2563eb; color:#1e40af; }
.alert-success { background:#d1fae5; border-left:4px solid #059669; color:#065f46; }
.alert-warning { background:#fef3c7; border-left:4px solid #d97706; color:#92400e; }
div[data-testid="stSidebarNav"] { display:none; }
.print-btn { background:#1a3a5c!important; color:white!important; }
</style>
""", unsafe_allow_html=True)

# ============================================================
# DATABASE
# ============================================================
def get_connection():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def hash_password(p): return hashlib.sha256(p.encode()).hexdigest()

def init_db():
    conn = get_connection(); c = conn.cursor()

    c.execute("""CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT UNIQUE NOT NULL,
        password TEXT NOT NULL, role TEXT NOT NULL, full_name TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)""")

    c.execute("""CREATE TABLE IF NOT EXISTS entete_cabinet (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nom_medecin TEXT, specialite TEXT, diplomes TEXT,
        cabinet TEXT, adresse TEXT, telephone TEXT,
        email TEXT, horaires TEXT, updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)""")

    c.execute("""CREATE TABLE IF NOT EXISTS patients (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nom TEXT NOT NULL, prenom TEXT NOT NULL,
        date_naissance TEXT, sexe TEXT, telephone TEXT,
        adresse TEXT, mutuelle TEXT, num_securite_sociale TEXT,
        antecedents TEXT, allergies TEXT, groupe_sanguin TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP, created_by TEXT)""")

    c.execute("""CREATE TABLE IF NOT EXISTS rendez_vous (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_id INTEGER, date TEXT NOT NULL, heure TEXT NOT NULL,
        motif TEXT, statut TEXT DEFAULT 'planifié', notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(patient_id) REFERENCES patients(id))""")

    c.execute("""CREATE TABLE IF NOT EXISTS consultations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_id INTEGER, date TEXT NOT NULL,
        anamnese TEXT, examen_clinique TEXT, diagnostic TEXT,
        traitement TEXT, notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(patient_id) REFERENCES patients(id))""")

    c.execute("""CREATE TABLE IF NOT EXISTS ordonnances (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_id INTEGER, consultation_id INTEGER, date TEXT NOT NULL,
        medicaments TEXT, instructions TEXT, duree TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(patient_id) REFERENCES patients(id))""")

    c.execute("""CREATE TABLE IF NOT EXISTS medicaments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nom TEXT NOT NULL, dci TEXT, classe TEXT, forme TEXT,
        dosage TEXT, posologie_adulte TEXT, posologie_enfant TEXT,
        contre_indications TEXT, effets_indesirables TEXT,
        prix REAL, stock INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)""")

    c.execute("""CREATE TABLE IF NOT EXISTS recettes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_id INTEGER, consultation_id INTEGER,
        date TEXT NOT NULL, acte TEXT, montant REAL,
        mode_paiement TEXT, paye INTEGER DEFAULT 0, notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(patient_id) REFERENCES patients(id))""")

    c.execute("""CREATE TABLE IF NOT EXISTS radios (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_id INTEGER, date TEXT NOT NULL,
        region TEXT, type_radio TEXT,
        fichier_original TEXT, fichier_traite TEXT,
        notes TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(patient_id) REFERENCES patients(id))""")

    conn.commit()

    # Défauts
    c.execute("SELECT COUNT(*) FROM users")
    if c.fetchone()[0] == 0:
        c.execute("INSERT INTO users (username,password,role,full_name) VALUES (?,?,?,?)",
                  ("medecin", hash_password("medecin123"), "medecin", "Dr. Bensalem Karim"))
        c.execute("INSERT INTO users (username,password,role,full_name) VALUES (?,?,?,?)",
                  ("secretaire", hash_password("secret123"), "secretaire", "Mme. Boudali Nadia"))

    c.execute("SELECT COUNT(*) FROM entete_cabinet")
    if c.fetchone()[0] == 0:
        c.execute("""INSERT INTO entete_cabinet
            (nom_medecin,specialite,diplomes,cabinet,adresse,telephone,email,horaires)
            VALUES (?,?,?,?,?,?,?,?)""",
            ("Dr. Bensalem Karim", "Traumatologie & Orthopédie",
             "Ex-Chef de Clinique CHU Tlemcen",
             "Cabinet Médical Dr. Bensalem",
             "12 Rue Ibn Khaldoun, Tlemcen 13000",
             "043 XX XX XX", "dr.bensalem@cabinet.dz",
             "Dim–Jeu : 08h00–17h00"))

    c.execute("SELECT COUNT(*) FROM medicaments")
    if c.fetchone()[0] == 0:
        meds = [
            ("Ibuprofène 400mg","Ibuprofène","AINS","Comprimé","400mg","400mg 3x/jour pendant les repas, max 1200mg/j","20-30mg/kg/j en 3 prises","Ulcère gastrique, insuffisance rénale, grossesse T3","Douleurs gastriques, nausées",150.0,50),
            ("Diclofénac 50mg","Diclofénac sodique","AINS","Comprimé","50mg","50mg 2-3x/jour, max 150mg/j","Non recommandé < 14 ans","Insuffisance cardiaque, rénale, ATCD ulcère","Douleurs abdominales, éruption",180.0,40),
            ("Kétoprofène 100mg","Kétoprofène","AINS","Comprimé LP","100mg","100-200mg/j en 1-2 prises, avec repas","Non recommandé < 15 ans","Allergie AINS, ulcère actif","Nausées, photosensibilité",200.0,30),
            ("Célécoxib 200mg","Célécoxib","AINS Cox-2","Gélule","200mg","200mg 1-2x/jour","Non recommandé enfant","Allergie sulfonamides, insuffisance hépatique","Douleurs abdominales, HTA",350.0,25),
            ("Paracétamol 1g","Paracétamol","Antalgique","Comprimé","1g","1g 3-4x/jour, max 4g/j, toutes les 6h min","15mg/kg toutes les 6h","Insuffisance hépatocellulaire","Rares si posologie respectée",80.0,100),
            ("Tramadol 50mg","Tramadol","Opioïde faible","Gélule","50mg","50-100mg 3-4x/jour, max 400mg/j","Non recommandé < 12 ans","Épilepsie non contrôlée, IMAOthérapie","Nausées, constipation, somnolence",220.0,30),
            ("Tramadol LP 100mg","Tramadol","Opioïde faible","Comprimé LP","100mg","100-200mg 2x/jour à 12h d'intervalle","Non recommandé < 18 ans","Épilepsie, dépression respiratoire","Somnolence, vertiges, nausées",280.0,20),
            ("Codéine + Paracétamol","Codéine/Paracétamol","Antalgique pallier 2","Comprimé","30/500mg","1-2 cp 3-4x/jour, max 6 cp/j","Non recommandé < 12 ans","Insuffisance hépatique, asthme, allaitement","Constipation, somnolence",190.0,25),
            ("Méthocarbamol 500mg","Méthocarbamol","Myorelaxant","Comprimé","500mg","1500mg 4x/jour 1ère semaine, puis 750mg 4x/j","Non établi < 16 ans","Épilepsie, myasthénie","Somnolence, vertiges",260.0,20),
            ("Thiocolchicoside 4mg","Thiocolchicoside","Myorelaxant","Comprimé","4mg","4mg 2x/jour, max 7 jours","Non recommandé < 16 ans","Grossesse, troubles convulsifs","Diarrhée, somnolence",300.0,20),
            ("Prednisone 5mg","Prednisone","Corticoïde","Comprimé","5mg","0.5-1mg/kg/j en cure courte","0.5-2mg/kg/j","Infections sévères non traitées","Prise de poids, HTA, hyperglycémie",120.0,40),
            ("Méthylprednisolone 16mg","Méthylprednisolone","Corticoïde","Comprimé","16mg","16-32mg/j en 1-2 prises matin","Variable","Idem prednisone","Idem corticoïdes",250.0,20),
            ("Énoxaparine 4000UI","Énoxaparine sodique","HBPM","Seringue SC","0.4mL","1 inj SC/jour prophylaxie; 100UI/kg 2x/j curatif","Adapté au poids","Hémorragie active, ATCD TIH","Hématome point injection",400.0,30),
            ("Rivaroxaban 10mg","Rivaroxaban","AOD","Comprimé","10mg","10mg 1x/jour après chirurgie orthopédique","Non recommandé < 18 ans","Hémorragie active, grossesse","Saignements, nausées",850.0,20),
            ("Alendronate 70mg","Alendronate sodique","Bisphosphonate","Comprimé","70mg","70mg 1x/semaine à jeun, rester debout 30min","Non utilisé enfant","Hypocalcémie, pathologies œsophagiennes","Douleurs osseuses, musculaires",320.0,15),
            ("Calcium + Vitamine D3","Calcium/Cholécalciférol","Supplément","Comprimé à croquer","1g/800UI","1-2 cp/jour pendant ou après repas","500mg/400UI 1x/j selon âge","Hypercalcémie, calculs rénaux","Constipation, flatulences",180.0,50),
            ("Oméprazole 20mg","Oméprazole","IPP","Gélule","20mg","20mg 1x/jour à jeun, à prendre avec AINS","0.7-1.5mg/kg/j","Allergie aux IPP","Céphalées, diarrhée, nausées",150.0,60),
            ("Pantoprazole 40mg","Pantoprazole","IPP","Comprimé","40mg","40mg 1x/jour à jeun","Non recommandé < 12 ans","Allergie benzimidazoles","Diarrhée, nausées, céphalées",160.0,50),
            ("Amoxicilline 1g","Amoxicilline","Antibiotique","Comprimé","1g","1g 3x/jour 7-10 jours","50mg/kg/j en 3 prises","Allergie pénicilline, mononucléose","Diarrhée, éruption, nausées",200.0,30),
            ("Amoxicilline/Ac.clavulanique 1g","Amoxicilline+Ac.clavulanique","Antibiotique","Comprimé","1g","1g 3x/jour 7-10 jours","25-45mg/kg/j en 3 prises","Allergie pénicilline, ictère cholestatique","Diarrhée, nausées, candidose",350.0,25),
            ("Ciprofloxacine 500mg","Ciprofloxacine","Fluoroquinolone","Comprimé","500mg","500-750mg 2x/jour 7-10 jours","Non recommandé < 18 ans","Grossesse, épilepsie, déficit G6PD","Tendinopathies, photosensibilité",280.0,20),
            ("Alpha-amylase 6000UI","Alpha-amylase","Antioedémateux","Comprimé","6000UI","3 cp 3x/jour phase aiguë, puis 2 cp 3x/j","Adapté","Allergie aux enzymes","Diarrhées légères",220.0,25),
            ("Diclofénac gel 1%","Diclofénac diéthylamine","AINS topique","Gel","1%","4g 3-4x/jour sur zone douloureuse, masser doucement","Non recommandé < 14 ans","Peau lésée, dermatose, grossesse T3","Prurit local, éruption",250.0,20),
            ("Kétoprofène gel 2.5%","Kétoprofène","AINS topique","Gel","2.5%","2g 1-2x/jour, éviter exposition soleil","Non recommandé < 15 ans","Exposition solaire, peau lésée","Photosensibilité, prurit",230.0,15),
        ]
        c.executemany("INSERT INTO medicaments (nom,dci,classe,forme,dosage,posologie_adulte,posologie_enfant,contre_indications,effets_indesirables,prix,stock) VALUES (?,?,?,?,?,?,?,?,?,?,?)", meds)
    conn.commit(); conn.close()

# ============================================================
# UTILITAIRES
# ============================================================
def get_entete():
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT * FROM entete_cabinet ORDER BY id DESC LIMIT 1")
    row = c.fetchone(); conn.close()
    return dict(row) if row else {}

def calc_age(dob_str):
    if not dob_str: return ""
    try:
        dob = datetime.strptime(dob_str, "%Y-%m-%d")
        return str((date.today() - dob.date()).days // 365)
    except: return ""

def set_cell_bg(cell, hex_color):
    """Couleur de fond d'une cellule tableau"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        if val is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), val.get('val','single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11, color=None, font="Arial"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)
        run.font.color.rgb = RGBColor(r, g, b)
    return run

def gen_ordonnance_docx(patient, meds_list, date_ordo, diagnostic="",
                         instructions_gen="", prochain_rdv="", arret_travail=""):
    """Génère un DOCX professionnel avec python-docx — aucune dépendance Node.js"""
    try:
        entete = get_entete()
        doc = Document()

        # ── Marges de page ─────────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin   = Cm(2.0)
            section.right_margin  = Cm(2.0)

        BLEU     = "1a3a5c"
        GRIS     = "f3f4f6"
        BLEU_CL  = "dbeafe"
        ROUGE    = "dc2626"
        ORANGE   = "92400e"
        GRIS_TEX = "6b7280"

        # ── EN-TÊTE : tableau 2 colonnes ───────────────────────
        t_header = doc.add_table(rows=1, cols=2)
        t_header.style = 'Table Grid'
        w_left  = t_header.cell(0,0)
        w_right = t_header.cell(0,1)
        # Largeurs
        w_left._tc.get_or_add_tcPr().append(
            OxmlElement('w:tcW'))
        for cell, w_twips in [(w_left, 5400), (w_right, 3400)]:
            tcW = cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW'); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn('w:w'), str(w_twips)); tcW.set(qn('w:type'), 'dxa')

        # Cellule gauche — infos médecin
        set_cell_borders(w_left)
        set_cell_borders(w_right)
        set_cell_bg(w_right, BLEU)

        p = w_left.paragraphs[0]
        add_run(p, entete.get('nom_medecin','Dr. Nom Prénom'), bold=True, size=14, color=BLEU)
        p2 = w_left.add_paragraph()
        add_run(p2, entete.get('specialite',''), size=11, color="374151")
        p3 = w_left.add_paragraph()
        add_run(p3, entete.get('diplomes',''), size=9, italic=True, color=GRIS_TEX)

        # Cellule droite — infos cabinet (fond bleu, texte blanc)
        pr = w_right.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(pr, entete.get('cabinet','Cabinet Médical'), bold=True, size=11, color="FFFFFF")
        for txt in [entete.get('adresse',''), 
                    ('Tél: '+entete.get('telephone','')) if entete.get('telephone') else '',
                    entete.get('email',''),
                    entete.get('horaires','')]:
            if txt:
                pp = w_right.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(pp, txt, size=9, color="FFFFFF")

        # Ligne séparatrice bleue
        doc.add_paragraph()
        sep = doc.add_paragraph()
        pPr = sep._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:color'), BLEU); pBdr.append(bottom); pPr.append(pBdr)

        # ── PATIENT + DATE : tableau 2 colonnes ────────────────
        t_pat = doc.add_table(rows=1, cols=2)
        t_pat.style = 'Table Grid'
        c_pat  = t_pat.cell(0,0)
        c_date = t_pat.cell(0,1)
        set_cell_borders(c_pat); set_cell_borders(c_date)
        set_cell_bg(c_pat, GRIS)

        pat_nom = f"{patient.get('prenom','')} {patient.get('nom','')}"
        age_str = calc_age(patient.get('date_naissance',''))

        pp = c_pat.paragraphs[0]
        add_run(pp, "PATIENT : ", bold=True, size=11, color=BLEU)
        add_run(pp, pat_nom, bold=True, size=11)
        pp2 = c_pat.add_paragraph()
        add_run(pp2, "Né(e) le : ", size=10, color=GRIS_TEX)
        add_run(pp2, (patient.get('date_naissance','') or '—') + f"   Âge : {age_str} ans" if age_str else (patient.get('date_naissance','') or '—'), size=10)
        if patient.get('mutuelle'):
            pm = c_pat.add_paragraph()
            add_run(pm, "Mutuelle : ", size=10, color=GRIS_TEX)
            add_run(pm, patient.get('mutuelle',''), size=10)
        if patient.get('allergies'):
            pa = c_pat.add_paragraph()
            add_run(pa, "⚠ ALLERGIES : ", bold=True, size=10, color=ROUGE)
            add_run(pa, patient.get('allergies',''), size=10, color=ROUGE)

        pd = c_date.paragraphs[0]; pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(pd, f"Le : {date_ordo}", bold=True, size=11)
        if diagnostic:
            pd2 = c_date.add_paragraph(); pd2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(pd2, f"Diag. : {diagnostic}", italic=True, size=9, color="374151")

        # ── TITRE ──────────────────────────────────────────────
        doc.add_paragraph()
        titre = doc.add_paragraph()
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(titre, "ORDONNANCE MÉDICALE", bold=True, size=14, color=BLEU)

        sep2 = doc.add_paragraph()
        pPr2 = sep2._p.get_or_add_pPr()
        pBdr2 = OxmlElement('w:pBdr')
        bot2 = OxmlElement('w:bottom')
        bot2.set(qn('w:val'), 'single'); bot2.set(qn('w:sz'), '4')
        bot2.set(qn('w:color'), BLEU_CL); pBdr2.append(bot2); pPr2.append(pBdr2)
        doc.add_paragraph()

        # ── MÉDICAMENTS ────────────────────────────────────────
        for i, med in enumerate(meds_list):
            # Numéro + Nom
            pm = doc.add_paragraph()
            add_run(pm, f"{i+1}.  ", bold=True, size=12, color=BLEU)
            add_run(pm, med.get('nom',''), bold=True, size=12)
            if med.get('dci'):
                add_run(pm, f"  ({med['dci']})", italic=True, size=10, color=GRIS_TEX)
            # Posologie
            pp_pos = doc.add_paragraph(); pp_pos.paragraph_format.left_indent = Cm(1)
            add_run(pp_pos, "Posologie : ", bold=True, size=10, color="374151")
            add_run(pp_pos, med.get('posologie','—'), size=10)
            # Durée
            pp_dur = doc.add_paragraph(); pp_dur.paragraph_format.left_indent = Cm(1)
            add_run(pp_dur, "Durée : ", bold=True, size=10, color="374151")
            add_run(pp_dur, med.get('duree','—'), size=10)
            # Instructions
            if med.get('instructions'):
                pp_inst = doc.add_paragraph(); pp_inst.paragraph_format.left_indent = Cm(1)
                add_run(pp_inst, "⚠ ", size=10)
                add_run(pp_inst, med['instructions'], italic=True, size=10, color=ORANGE)
            # Séparateur entre médicaments
            if i < len(meds_list) - 1:
                sep_m = doc.add_paragraph()
                pPrm = sep_m._p.get_or_add_pPr()
                pBdrm = OxmlElement('w:pBdr')
                botm = OxmlElement('w:bottom')
                botm.set(qn('w:val'), 'dashed'); botm.set(qn('w:sz'), '4')
                botm.set(qn('w:color'), 'e5e7eb'); pBdrm.append(botm); pPrm.append(pBdrm)

        # ── INSTRUCTIONS GÉNÉRALES ──────────────────────────────
        if instructions_gen:
            doc.add_paragraph()
            pi = doc.add_paragraph()
            add_run(pi, "Instructions générales : ", bold=True, size=10, color=BLEU)
            add_run(pi, instructions_gen, italic=True, size=10)

        # ── SIGNATURE + RDV ────────────────────────────────────
        doc.add_paragraph(); doc.add_paragraph()
        t_sig = doc.add_table(rows=1, cols=2)
        t_sig.style = 'Table Grid'
        c_rdv = t_sig.cell(0,0); c_sign = t_sig.cell(0,1)
        set_cell_borders(c_rdv); set_cell_borders(c_sign)

        if prochain_rdv:
            pr_p = c_rdv.paragraphs[0]
            add_run(pr_p, "Prochain RDV : ", bold=True, size=10, color=BLEU)
            add_run(pr_p, prochain_rdv, size=10)
        if arret_travail:
            pa_p = c_rdv.add_paragraph()
            add_run(pa_p, "Arrêt de travail : ", bold=True, size=10, color=ROUGE)
            add_run(pa_p, arret_travail, size=10, color=ROUGE)

        ps = c_sign.paragraphs[0]; ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(ps, entete.get('nom_medecin',''), bold=True, size=10, color=BLEU)
        ps2 = c_sign.add_paragraph(); ps2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(ps2, entete.get('specialite',''), italic=True, size=9, color=GRIS_TEX)
        # Espace signature
        for _ in range(4): c_sign.add_paragraph()
        ps_final = c_sign.add_paragraph(); ps_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr_sig = ps_final._p.get_or_add_pPr()
        pBdr_sig = OxmlElement('w:pBdr')
        top_sig = OxmlElement('w:top')
        top_sig.set(qn('w:val'), 'single'); top_sig.set(qn('w:sz'), '4')
        top_sig.set(qn('w:color'), '374151'); pBdr_sig.append(top_sig); pPr_sig.append(pBdr_sig)
        add_run(ps_final, "Cachet & Signature", size=9, color="9ca3af")

        # ── PIED DE PAGE ───────────────────────────────────────
        doc.add_paragraph()
        pied = doc.add_paragraph()
        pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr_p = pied._p.get_or_add_pPr()
        pBdr_p = OxmlElement('w:pBdr')
        top_p = OxmlElement('w:top')
        top_p.set(qn('w:val'), 'single'); top_p.set(qn('w:sz'), '4')
        top_p.set(qn('w:color'), BLEU_CL); pBdr_p.append(top_p); pPr_p.append(pBdr_p)
        footer_txt = " | ".join(filter(None, [
            entete.get('cabinet',''), entete.get('adresse',''),
            entete.get('telephone',''), entete.get('horaires','')]))
        add_run(pied, footer_txt, size=8, color="9ca3af")

        # ── Sauvegarder en mémoire ─────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue(), None

    except Exception as e:
        return None, str(e)

def save_radio_to_patient(patient_id, img, region, type_radio, notes, is_traited=False):
    """Sauvegarde une image radio dans le dossier du patient"""
    patient_radio_dir = os.path.join(RADIO_DIR, str(patient_id))
    os.makedirs(patient_radio_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    suffix = "_traite" if is_traited else "_original"
    fname = f"radio_{ts}{suffix}.png"
    fpath = os.path.join(patient_radio_dir, fname)
    img.save(fpath, format="PNG", quality=95)
    return fname

def backup_all_data():
    """Crée un ZIP complet de toutes les données"""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_path = os.path.join(EXPORT_DIR, f"backup_cabinet_{ts}.zip")
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.write(DB_PATH, "cabinet.db")
        if os.path.exists(RADIO_DIR):
            for root, dirs, files in os.walk(RADIO_DIR):
                for file in files:
                    fp = os.path.join(root, file)
                    zf.write(fp, os.path.relpath(fp, os.path.join(RADIO_DIR, '..')))
    with open(zip_path, 'rb') as f:
        return f.read(), f"backup_cabinet_{ts}.zip"

# ============================================================
# AUTH
# ============================================================
def page_login():
    st.markdown("""<div style='text-align:center;margin-bottom:2rem;'>
        <div style='font-size:3rem;'>🦴</div>
        <h2 style='color:#1a3a5c;margin:0;'>Cabinet Médical</h2>
        <p style='color:#6b7280;font-size:0.9rem;'>Traumatologie & Orthopédie</p>
    </div>""", unsafe_allow_html=True)
    col1,col2,col3 = st.columns([1,2,1])
    with col2:
        st.markdown("### 🔐 Connexion")
        username = st.text_input("Identifiant")
        password = st.text_input("Mot de passe", type="password")
        if st.button("Se connecter", use_container_width=True, type="primary"):
            if username and password:
                conn = get_connection(); c = conn.cursor()
                c.execute("SELECT * FROM users WHERE username=? AND password=?",
                          (username, hash_password(password)))
                user = c.fetchone(); conn.close()
                if user:
                    st.session_state.logged_in = True
                    st.session_state.role      = user["role"]
                    st.session_state.username  = username
                    st.session_state.full_name = user["full_name"]
                    st.rerun()
                else: st.error("Identifiant ou mot de passe incorrect")
            else: st.warning("Veuillez remplir tous les champs")
        st.markdown("<div style='text-align:center;margin-top:1rem;font-size:0.78rem;color:#9ca3af;'><b>Médecin:</b> medecin / medecin123 &nbsp;|&nbsp; <b>Secrétaire:</b> secretaire / secret123</div>", unsafe_allow_html=True)

# ============================================================
# PAGE : ACCUEIL
# ============================================================
def page_accueil():
    role = st.session_state.role; name = st.session_state.full_name
    st.markdown(f"""<div class='main-header'><div style='font-size:2rem;'>🦴</div>
        <div><h1>{'Tableau de Bord' if role=='medecin' else 'Accueil Secrétariat'}</h1>
        <p>Bonjour, {name} · {date.today().strftime('%A %d %B %Y')}</p></div>
        <div class='role-badge'>{'👨‍⚕️ Médecin' if role=='medecin' else '👩‍💼 Secrétaire'}</div></div>""", unsafe_allow_html=True)
    conn = get_connection(); c = conn.cursor(); today = date.today().isoformat()
    c.execute("SELECT COUNT(*) FROM patients"); tp = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM rendez_vous WHERE date=?", (today,)); rt = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM rendez_vous WHERE statut='planifié' AND date>=?", (today,)); rp = c.fetchone()[0]
    c.execute("SELECT COALESCE(SUM(montant),0) FROM recettes WHERE date=? AND paye=1", (today,)); rec = c.fetchone()[0]
    col1,col2,col3,col4 = st.columns(4)
    with col1: st.markdown(f"<div class='stat-card'><div class='number'>{tp}</div><div class='label'>👤 Patients</div></div>", unsafe_allow_html=True)
    with col2: st.markdown(f"<div class='stat-card green'><div class='number'>{rt}</div><div class='label'>📅 RDV aujourd'hui</div></div>", unsafe_allow_html=True)
    with col3: st.markdown(f"<div class='stat-card orange'><div class='number'>{rp}</div><div class='label'>⏳ RDV à venir</div></div>", unsafe_allow_html=True)
    with col4: st.markdown(f"<div class='stat-card purple'><div class='number'>{rec:,.0f} DA</div><div class='label'>💰 Recette du jour</div></div>", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    cl,cr = st.columns([3,2])
    with cl:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>📅 Rendez-vous du jour</div>", unsafe_allow_html=True)
        c.execute("SELECT r.*,p.nom,p.prenom,p.telephone FROM rendez_vous r JOIN patients p ON r.patient_id=p.id WHERE r.date=? ORDER BY r.heure", (today,))
        for rdv in c.fetchall():
            sc={"planifié":"badge-blue","en cours":"badge-orange","terminé":"badge-green","annulé":"badge-red"}.get(rdv["statut"],"badge-gray")
            st.markdown(f"<div class='rdv-card'><b>🕐 {rdv['heure']}</b> — <b>{rdv['prenom']} {rdv['nom']}</b> <span class='badge {sc}'>{rdv['statut']}</span><br><span style='font-size:0.82rem;color:#374151;'>📌 {rdv['motif'] or '—'} | 📞 {rdv['telephone'] or '—'}</span></div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    with cr:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>🆕 Derniers patients</div>", unsafe_allow_html=True)
        c.execute("SELECT nom,prenom,date_naissance FROM patients ORDER BY created_at DESC LIMIT 6")
        for p in c.fetchall():
            st.markdown(f"<div style='padding:0.5rem 0;border-bottom:1px solid #f3f4f6;'><b>👤 {p['prenom']} {p['nom']}</b><br><span style='font-size:0.78rem;color:#6b7280;'>{p['date_naissance'] or '—'}</span></div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
        if role == "medecin":
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown("<div class='section-title'>💊 Stock faible</div>", unsafe_allow_html=True)
            c.execute("SELECT nom,stock FROM medicaments WHERE stock<10 ORDER BY stock LIMIT 5")
            faibles = c.fetchall()
            if faibles:
                for m in faibles:
                    col = "#fee2e2" if m["stock"]<5 else "#fef3c7"
                    st.markdown(f"<div style='background:{col};padding:0.4rem 0.8rem;border-radius:6px;margin-bottom:0.3rem;font-size:0.85rem;'>💊 {m['nom']} — <b>{m['stock']} unités</b></div>", unsafe_allow_html=True)
            else:
                st.markdown("<div class='alert-success alert-box'>✅ Stocks suffisants</div>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
    conn.close()

# ============================================================
# PAGE : PARAMÈTRES EN-TÊTE
# ============================================================
def page_entete():
    st.markdown("<div class='main-header'><div style='font-size:2rem;'>⚙️</div><div><h1>Paramètres de l'En-tête</h1><p>Personnalisez l'en-tête de vos ordonnances</p></div></div>", unsafe_allow_html=True)
    conn = get_connection(); c = conn.cursor()
    c.execute("SELECT * FROM entete_cabinet ORDER BY id DESC LIMIT 1")
    row = c.fetchone(); h = dict(row) if row else {}

    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.markdown("<div class='section-title'>👨‍⚕️ Informations du médecin</div>", unsafe_allow_html=True)
    col1,col2 = st.columns(2)
    with col1:
        nom_med  = st.text_input("Nom du médecin", value=h.get("nom_medecin",""))
        spec     = st.text_input("Spécialité", value=h.get("specialite",""))
        diplomes = st.text_input("Diplômes / Titres", value=h.get("diplomes",""))
    with col2:
        cabinet  = st.text_input("Nom du cabinet", value=h.get("cabinet",""))
        adresse  = st.text_area("Adresse complète", value=h.get("adresse",""), height=68)
        tel      = st.text_input("Téléphone", value=h.get("telephone",""))
    col3,col4 = st.columns(2)
    with col3: email    = st.text_input("Email", value=h.get("email",""))
    with col4: horaires = st.text_input("Horaires de consultation", value=h.get("horaires",""))

    # Aperçu
    st.markdown("---")
    st.markdown("<div class='section-title'>👁️ Aperçu de l'en-tête</div>", unsafe_allow_html=True)
    st.markdown(f"""
    <div style='border:2px solid #1a3a5c;border-radius:8px;overflow:hidden;'>
        <div style='display:flex;'>
            <div style='flex:1;padding:1rem;'>
                <div style='font-size:1.2rem;font-weight:700;color:#1a3a5c;'>{nom_med or "Dr. Nom Prénom"}</div>
                <div style='color:#374151;font-size:0.9rem;'>{spec or "Spécialité"}</div>
                <div style='color:#6b7280;font-size:0.8rem;font-style:italic;'>{diplomes or "Diplômes et titres"}</div>
            </div>
            <div style='background:#1a3a5c;color:white;padding:1rem;min-width:200px;text-align:center;'>
                <div style='font-weight:600;font-size:0.9rem;'>{cabinet or "Nom du cabinet"}</div>
                <div style='font-size:0.78rem;opacity:0.9;margin-top:0.3rem;'>{adresse or "Adresse"}</div>
                <div style='font-size:0.78rem;opacity:0.9;'>📞 {tel or "Téléphone"}</div>
                <div style='font-size:0.75rem;opacity:0.8;'>{email or ""}</div>
            </div>
        </div>
        <div style='background:#1a3a5c;height:4px;'></div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("💾 Sauvegarder l'en-tête", type="primary", use_container_width=True):
        if h:
            c.execute("""UPDATE entete_cabinet SET nom_medecin=?,specialite=?,diplomes=?,
                cabinet=?,adresse=?,telephone=?,email=?,horaires=?,updated_at=CURRENT_TIMESTAMP
                WHERE id=?""",
                (nom_med,spec,diplomes,cabinet,adresse,tel,email,horaires,h["id"]))
        else:
            c.execute("INSERT INTO entete_cabinet (nom_medecin,specialite,diplomes,cabinet,adresse,telephone,email,horaires) VALUES (?,?,?,?,?,?,?,?)",
                      (nom_med,spec,diplomes,cabinet,adresse,tel,email,horaires))
        conn.commit()
        st.success("✅ En-tête sauvegardé ! Vos prochaines ordonnances utiliseront ces informations.")
    st.markdown("</div>", unsafe_allow_html=True)

    # Test génération
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.markdown("<div class='section-title'>🧪 Tester l'ordonnance</div>", unsafe_allow_html=True)
    st.info("Générez une ordonnance de test pour vérifier le rendu final.")
    if st.button("📄 Générer ordonnance de test"):
        test_patient = {"nom":"PATIENT","prenom":"Test","date_naissance":"1980-01-15","mutuelle":"CNAS","allergies":""}
        test_meds = [
            {"nom":"Ibuprofène 400mg","dci":"Ibuprofène","posologie":"400mg 3x/jour avec repas","duree":"5 jours","instructions":"Gastroprotection recommandée"},
            {"nom":"Oméprazole 20mg","dci":"Oméprazole","posologie":"20mg 1x/jour à jeun","duree":"5 jours","instructions":""}
        ]
        with st.spinner("Génération en cours..."):
            docx_bytes, err = gen_ordonnance_docx(test_patient, test_meds, date.today().strftime("%d/%m/%Y"), "Test - Lombalgie aiguë", "Repos relatif recommandé.")
        if docx_bytes:
            st.success("✅ Ordonnance générée avec succès !")
            st.download_button("⬇️ Télécharger l'ordonnance test (.docx)",
                docx_bytes, "ordonnance_test.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        else:
            st.error(f"❌ Erreur: {err}")
    st.markdown("</div>", unsafe_allow_html=True)
    conn.close()

# ============================================================
# PAGE : PATIENTS
# ============================================================
def page_patients():
    st.markdown("<div class='main-header'><div style='font-size:2rem;'>👤</div><div><h1>Gestion des Patients</h1><p>Dossiers et suivi des patients</p></div></div>", unsafe_allow_html=True)
    tab1,tab2,tab3 = st.tabs(["📋 Liste","➕ Nouveau patient","🔍 Dossier complet"])
    conn = get_connection(); c = conn.cursor()

    with tab1:
        cs,cf = st.columns([3,1])
        with cs: search = st.text_input("🔍 Rechercher (nom, prénom, téléphone)")
        with cf: sf = st.selectbox("Sexe", ["Tous","Masculin","Féminin"])
        q="SELECT * FROM patients WHERE 1=1"; p=[]
        if search: q+=" AND (nom LIKE ? OR prenom LIKE ? OR telephone LIKE ?)"; p+=[f"%{search}%"]*3
        if sf!="Tous": q+=" AND sexe=?"; p.append(sf)
        c.execute(q+" ORDER BY nom", p); patients=c.fetchall()
        st.markdown(f"<p style='color:#6b7280;font-size:0.85rem;'>{len(patients)} patient(s)</p>", unsafe_allow_html=True)
        for pat in patients:
            with st.expander(f"👤 {pat['prenom']} {pat['nom']} — {pat['telephone'] or 'N° non renseigné'}"):
                c1,c2,c3=st.columns(3)
                with c1:
                    st.write(f"**Naissance:** {pat['date_naissance'] or '—'}")
                    st.write(f"**Sexe:** {pat['sexe'] or '—'}")
                    st.write(f"**Groupe sanguin:** {pat['groupe_sanguin'] or '—'}")
                with c2:
                    st.write(f"**Téléphone:** {pat['telephone'] or '—'}")
                    st.write(f"**Mutuelle:** {pat['mutuelle'] or '—'}")
                with c3:
                    st.write(f"**Allergies:** {pat['allergies'] or 'Aucune connue'}")
                    st.write(f"**Antécédents:** {pat['antecedents'] or '—'}")
                b1,b2,b3=st.columns(3)
                with b1:
                    if st.button("📅 RDV", key=f"rdv_{pat['id']}"):
                        st.session_state.rdv_patient_id=pat['id']; st.session_state.current_page="rendez_vous"; st.rerun()
                with b2:
                    if st.button("🔍 Dossier", key=f"dos_{pat['id']}"):
                        st.session_state.patient_selectionne=pat['id']; st.session_state.current_page="patients"; st.rerun()
                with b3:
                    if st.session_state.role=="medecin":
                        if st.button("📋 Consulter", key=f"cons_{pat['id']}"):
                            st.session_state.consult_patient_id=pat['id']; st.session_state.current_page="ordonnances"; st.rerun()

    with tab2:
        c1,c2=st.columns(2)
        with c1:
            nom=st.text_input("Nom *"); prenom=st.text_input("Prénom *")
            dob=st.date_input("Date de naissance",value=None,min_value=date(1920,1,1),max_value=date.today())
            sexe=st.selectbox("Sexe",["","Masculin","Féminin"])
            gs=st.selectbox("Groupe sanguin",["","A+","A-","B+","B-","AB+","AB-","O+","O-"])
        with c2:
            tel=st.text_input("Téléphone"); adr=st.text_area("Adresse",height=68)
            mut=st.text_input("Mutuelle"); nss=st.text_input("N° Sécurité Sociale")
        ant=st.text_area("Antécédents médicaux",height=80); all_=st.text_area("Allergies",height=68)
        if st.button("💾 Enregistrer", type="primary"):
            if nom and prenom:
                c.execute("INSERT INTO patients (nom,prenom,date_naissance,sexe,telephone,adresse,mutuelle,num_securite_sociale,antecedents,allergies,groupe_sanguin,created_by) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                    (nom.upper(),prenom,str(dob) if dob else None,sexe or None,tel,adr,mut,nss,ant,all_,gs or None,st.session_state.username))
                conn.commit(); st.success(f"✅ {prenom} {nom.upper()} enregistré !"); st.rerun()
            else: st.error("Nom et prénom obligatoires.")

    with tab3:
        c.execute("SELECT id,nom,prenom FROM patients ORDER BY nom"); all_p=c.fetchall()
        if not all_p: st.info("Aucun patient."); conn.close(); return
        opts={f"{p['prenom']} {p['nom']}":p['id'] for p in all_p}
        def_idx=0
        if "patient_selectionne" in st.session_state:
            for i,(k,v) in enumerate(opts.items()):
                if v==st.session_state.patient_selectionne: def_idx=i; break
        sel=st.selectbox("Patient",list(opts.keys()),index=def_idx); pid=opts[sel]
        c.execute("SELECT * FROM patients WHERE id=?",(pid,)); pat=c.fetchone()
        if pat:
            col1,col2=st.columns([2,1])
            with col1:
                st.markdown(f"""<div class='card'><div class='section-title'>🪪 Informations</div>
                    <table style='width:100%;'>
                    <tr><td style='color:#6b7280;'>Nom complet</td><td><b>{pat['prenom']} {pat['nom']}</b></td></tr>
                    <tr><td style='color:#6b7280;'>Naissance</td><td>{pat['date_naissance'] or '—'} ({calc_age(pat['date_naissance'])} ans)</td></tr>
                    <tr><td style='color:#6b7280;'>Groupe sanguin</td><td><b style='color:#dc2626;'>{pat['groupe_sanguin'] or '—'}</b></td></tr>
                    <tr><td style='color:#6b7280;'>Téléphone</td><td>{pat['telephone'] or '—'}</td></tr>
                    <tr><td style='color:#6b7280;'>Mutuelle</td><td>{pat['mutuelle'] or '—'}</td></tr>
                    </table></div>""", unsafe_allow_html=True)
            with col2:
                st.markdown(f"""<div class='card' style='background:#fef2f2;border-color:#fecaca;'>
                    <div class='section-title' style='color:#dc2626;'>⚠️ Alertes</div>
                    <p style='font-size:0.85rem;'><b>Allergies:</b><br>{pat['allergies'] or 'Aucune'}</p>
                    <p style='font-size:0.85rem;'><b>Antécédents:</b><br>{pat['antecedents'] or '—'}</p>
                </div>""", unsafe_allow_html=True)
            if st.session_state.role=="medecin":
                st.markdown("---")
                t_cons, t_radio, t_ordo = st.tabs(["📋 Consultations","🩻 Radios","💊 Ordonnances"])
                with t_cons:
                    c.execute("SELECT * FROM consultations WHERE patient_id=? ORDER BY date DESC",(pid,))
                    for cons in c.fetchall():
                        with st.expander(f"📋 {cons['date']} — {cons['diagnostic'] or '—'}"):
                            st.write(f"**Anamnèse:** {cons['anamnese'] or '—'}")
                            st.write(f"**Examen:** {cons['examen_clinique'] or '—'}")
                            st.write(f"**Diagnostic:** {cons['diagnostic'] or '—'}")
                with t_radio:
                    c.execute("SELECT * FROM radios WHERE patient_id=? ORDER BY date DESC",(pid,))
                    radios_pat = c.fetchall()
                    if radios_pat:
                        for rad in radios_pat:
                            with st.expander(f"🩻 {rad['date']} — {rad['region']} ({rad['type_radio']})"):
                                col_o, col_t = st.columns(2)
                                patient_radio_dir = os.path.join(RADIO_DIR, str(pid))
                                with col_o:
                                    st.write("**Image originale:**")
                                    if rad['fichier_original']:
                                        fp = os.path.join(patient_radio_dir, rad['fichier_original'])
                                        if os.path.exists(fp):
                                            st.image(fp, use_container_width=True)
                                            with open(fp,'rb') as f:
                                                st.download_button("⬇️ Télécharger originale", f.read(), rad['fichier_original'], key=f"dl_o_{rad['id']}")
                                with col_t:
                                    st.write("**Image traitée:**")
                                    if rad['fichier_traite']:
                                        fp2 = os.path.join(patient_radio_dir, rad['fichier_traite'])
                                        if os.path.exists(fp2):
                                            st.image(fp2, use_container_width=True)
                                            with open(fp2,'rb') as f:
                                                st.download_button("⬇️ Télécharger traitée", f.read(), rad['fichier_traite'], key=f"dl_t_{rad['id']}")
                                if rad['notes']:
                                    st.info(f"📝 {rad['notes']}")
                    else:
                        st.info("Aucune radio enregistrée pour ce patient.")
                with t_ordo:
                    c.execute("SELECT * FROM ordonnances WHERE patient_id=? ORDER BY date DESC",(pid,))
                    for o in c.fetchall():
                        with st.expander(f"💊 {o['date']}"):
                            for line in (o['medicaments'] or '').split('\n'):
                                if line.strip(): st.write(f"• {line}")
    conn.close()

# ============================================================
# PAGE : RENDEZ-VOUS
# ============================================================
MOTIFS_RDV = ["Consultation initiale","Fracture / Traumatisme","Entorse / Luxation",
    "Lombalgie / Dorsalgie","Cervicalgie","Gonarthrose / Coxarthrose","Hernie discale",
    "Suivi post-opératoire","Contrôle radiologique","Tendinopathie","Ostéoporose","Certificat médical","Autre motif"]

def page_rendez_vous():
    st.markdown("<div class='main-header'><div style='font-size:2rem;'>📅</div><div><h1>Rendez-vous</h1><p>Agenda et planning du cabinet</p></div></div>", unsafe_allow_html=True)
    tab1,tab2,tab3=st.tabs(["📅 Agenda du jour","📆 Semaine","➕ Nouveau RDV"])
    conn=get_connection(); c=conn.cursor()

    with tab1:
        sel_date=st.date_input("📅 Date", value=date.today())
        c.execute("SELECT r.*,p.nom,p.prenom,p.telephone FROM rendez_vous r JOIN patients p ON r.patient_id=p.id WHERE r.date=? ORDER BY r.heure",(str(sel_date),))
        rdvs=c.fetchall()
        if rdvs:
            for rdv in rdvs:
                sc={"planifié":"badge-blue","en cours":"badge-orange","terminé":"badge-green","annulé":"badge-red"}.get(rdv["statut"],"badge-gray")
                cr,cs=st.columns([4,1])
                with cr: st.markdown(f"<div class='rdv-card'><b>🕐 {rdv['heure']}</b> — <b>{rdv['prenom']} {rdv['nom']}</b> <span class='badge {sc}'>{rdv['statut']}</span><br><span style='font-size:0.8rem;color:#374151;'>📌 {rdv['motif'] or '—'} | 📞 {rdv['telephone'] or '—'}</span></div>", unsafe_allow_html=True)
                with cs:
                    ns=st.selectbox("",["planifié","en cours","terminé","annulé"],index=["planifié","en cours","terminé","annulé"].index(rdv['statut']),key=f"st_{rdv['id']}")
                    if ns!=rdv['statut']:
                        if st.button("✅",key=f"upd_{rdv['id']}"):
                            c.execute("UPDATE rendez_vous SET statut=? WHERE id=?",(ns,rdv['id'])); conn.commit(); st.rerun()
        else: st.markdown("<div class='alert-info alert-box'>Aucun rendez-vous pour cette date.</div>", unsafe_allow_html=True)

    with tab2:
        today_dt=date.today(); ws=today_dt-timedelta(days=today_dt.weekday())
        days=[ws+timedelta(days=i) for i in range(6)]
        cols=st.columns(6)
        for col,day in zip(cols,days):
            c.execute("SELECT COUNT(*) FROM rendez_vous WHERE date=? AND statut!='annulé'",(str(day),)); cnt=c.fetchone()[0]
            bg="#1a3a5c" if day==today_dt else "#f8fafc"; cl="white" if day==today_dt else "#1a3a5c"
            with col: st.markdown(f"<div style='background:{bg};color:{cl};padding:0.8rem;border-radius:8px;text-align:center;border:1px solid #e5e7eb;'><div style='font-size:0.7rem;font-weight:600;'>{day.strftime('%A')[:3].upper()}</div><div style='font-size:1.4rem;font-weight:700;'>{day.day}</div><div style='font-size:0.72rem;'>{cnt} RDV</div></div>", unsafe_allow_html=True)
        st.markdown("<br>",unsafe_allow_html=True)
        for day in days:
            c.execute("SELECT r.heure,r.motif,p.nom,p.prenom FROM rendez_vous r JOIN patients p ON r.patient_id=p.id WHERE r.date=? AND r.statut!='annulé' ORDER BY r.heure",(str(day),))
            drdvs=c.fetchall()
            if drdvs:
                with st.expander(f"📅 {day.strftime('%A %d/%m')} — {len(drdvs)} RDV"):
                    for r in drdvs: st.write(f"• **{r['heure']}** — {r['prenom']} {r['nom']} _{r['motif'] or '—'}_")

    with tab3:
        c.execute("SELECT id,nom,prenom FROM patients ORDER BY nom"); all_p=c.fetchall()
        if not all_p: st.warning("⚠️ Aucun patient enregistré."); conn.close(); return
        opts={f"{p['prenom']} {p['nom']}":p['id'] for p in all_p}
        def_idx=0
        if "rdv_patient_id" in st.session_state:
            for i,(k,v) in enumerate(opts.items()):
                if v==st.session_state.rdv_patient_id: def_idx=i; break
            del st.session_state.rdv_patient_id
        c1,c2=st.columns(2)
        with c1:
            sel_p=st.selectbox("👤 Patient *",list(opts.keys()),index=def_idx)
            rdv_dt=st.date_input("📅 Date *",value=date.today(),min_value=date.today())
            motif=st.selectbox("📌 Motif *",MOTIFS_RDV)
        with c2:
            heures=[f"{h:02d}:{m:02d}" for h in range(8,18) for m in [0,15,30,45]]
            heure=st.selectbox("🕐 Heure *",heures,index=heures.index("09:00"))
            statut=st.selectbox("Statut",["planifié","en cours","terminé","annulé"])
            notes=st.text_area("Notes",height=80)
        if st.button("💾 Enregistrer le RDV",type="primary"):
            c.execute("SELECT id FROM rendez_vous WHERE date=? AND heure=? AND statut!='annulé'",(str(rdv_dt),heure))
            if c.fetchone(): st.error(f"⚠️ Créneau {heure} déjà pris.")
            else:
                c.execute("INSERT INTO rendez_vous (patient_id,date,heure,motif,statut,notes) VALUES (?,?,?,?,?,?)",(opts[sel_p],str(rdv_dt),heure,motif,statut,notes))
                conn.commit(); st.success(f"✅ RDV enregistré pour {sel_p} à {heure}"); st.rerun()
    conn.close()

# ============================================================
# PAGE : RADIOLOGIE (AMÉLIORÉE)
# ============================================================
def apply_clahe_manual(img_array, clip_limit=20, tile_size=32):
    def clahe_ch(ch, cl, ts):
        h,w=ch.shape; res=np.zeros_like(ch,dtype=np.float64)
        for i in range(0,h,ts):
            for j in range(0,w,ts):
                tile=ch[i:i+ts,j:j+ts]; hist,_=np.histogram(tile.flatten(),256,[0,256])
                excess=np.sum(np.maximum(hist-cl*hist.mean(),0)); hist=np.minimum(hist,cl*hist.mean()); hist+=excess/256
                cdf=hist.cumsum(); cdf_min=cdf[cdf>0].min() if cdf[cdf>0].size>0 else 0
                cdf_n=((cdf-cdf_min)/(tile.size-cdf_min+1e-7)*255).astype(np.uint8); res[i:i+ts,j:j+ts]=cdf_n[tile]
        return res.astype(np.uint8)
    if len(img_array.shape)==3: return np.stack([clahe_ch(img_array[:,:,k],clip_limit,tile_size) for k in range(3)],axis=2)
    return clahe_ch(img_array,clip_limit,tile_size)

def process_radio(img, p):
    if p.get("mode_radio"): img=img.convert("L")
    arr=np.array(img)
    if p.get("clahe"):
        try:
            import cv2
            if len(arr.shape)==3:
                lab=cv2.cvtColor(arr,cv2.COLOR_RGB2LAB); cl=cv2.createCLAHE(clipLimit=p.get("clahe_clip",2.0),tileGridSize=(8,8))
                lab[:,:,0]=cl.apply(lab[:,:,0]); arr=cv2.cvtColor(lab,cv2.COLOR_LAB2RGB)
            else:
                cl=cv2.createCLAHE(clipLimit=p.get("clahe_clip",2.0),tileGridSize=(8,8)); arr=cl.apply(arr)
        except ImportError: arr=apply_clahe_manual(arr,clip_limit=p.get("clahe_clip",2.0)*10,tile_size=32)
    img=Image.fromarray(arr)
    if p.get("contraste",1.0)!=1.0: img=ImageEnhance.Contrast(img).enhance(p["contraste"])
    if p.get("luminosite",1.0)!=1.0: img=ImageEnhance.Brightness(img).enhance(p["luminosite"])
    if p.get("nettete",1.0)!=1.0: img=ImageEnhance.Sharpness(img).enhance(p["nettete"])
    if p.get("debruitage"): img=img.filter(ImageFilter.GaussianBlur(radius=p.get("denoise_r",0.8)))
    if p.get("detail_osseux"): img=img.filter(ImageFilter.UnsharpMask(radius=2,percent=150,threshold=3))
    if p.get("egalisation"):
        if img.mode=="L": img=ImageOps.equalize(img)
        else:
            r,g,b=img.split(); img=Image.merge("RGB",(ImageOps.equalize(r),ImageOps.equalize(g),ImageOps.equalize(b)))
    if p.get("negatif"): img=ImageOps.invert(img.convert("RGB")); img=img.convert("L") if p.get("mode_radio") else img
    gamma=p.get("gamma",1.0)
    if gamma!=1.0:
        lut=[int(((i/255.0)**(1.0/gamma))*255) for i in range(256)]
        img=img.point(lut) if img.mode=="L" else img.point(lut*3)
    return img

def page_radiologie():
    st.markdown("<div class='main-header'><div style='font-size:2rem;'>🩻</div><div><h1>Traitement Radiologique</h1><p>Amélioration des clichés • Sauvegarde dans le dossier patient • Impression</p></div></div>", unsafe_allow_html=True)

    conn=get_connection(); c=conn.cursor()
    c.execute("SELECT id,nom,prenom FROM patients ORDER BY nom"); all_p=c.fetchall()
    opts_p={f"{p['prenom']} {p['nom']}":p['id'] for p in all_p} if all_p else {}

    col_l,col_r=st.columns([1,2])

    with col_l:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>⚙️ Paramètres</div>", unsafe_allow_html=True)
        uf=st.file_uploader("📁 Charger un cliché radiologique", type=["jpg","jpeg","png","tif","tiff","bmp"])
        if uf:
            mode_radio=st.checkbox("🩻 Mode radiologie (niveaux de gris)",value=True)
            use_clahe=st.checkbox("CLAHE (contraste adaptatif)",value=True)
            use_denoise=st.checkbox("Débruitage",value=False)
            use_detail=st.checkbox("Détail osseux",value=False)
            use_egal=st.checkbox("Égalisation histogramme",value=False)
            use_neg=st.checkbox("Négatif / Inversé",value=False)
            contraste=st.slider("Contraste",0.5,3.0,1.2,0.1)
            luminosite=st.slider("Luminosité",0.5,2.0,1.0,0.1)
            nettete=st.slider("Netteté",0.5,3.0,1.3,0.1)
            gamma=st.slider("Gamma",0.3,2.5,1.0,0.1)
            clahe_clip=st.slider("Clip CLAHE",1.0,5.0,2.0,0.5) if use_clahe else 2.0
            denoise_r=st.slider("Force débruitage",0.3,3.0,0.8,0.1) if use_denoise else 0.8
            params={"mode_radio":mode_radio,"clahe":use_clahe,"clahe_clip":clahe_clip,
                    "contraste":contraste,"luminosite":luminosite,"nettete":nettete,"gamma":gamma,
                    "debruitage":use_denoise,"denoise_r":denoise_r,"detail_osseux":use_detail,
                    "negatif":use_neg,"egalisation":use_egal}
            cb1,cb2=st.columns(2)
            with cb1:
                if st.button("🚀 Traiter",type="primary",use_container_width=True):
                    st.session_state.radio_img=process_radio(Image.open(uf),params)
                    st.session_state.radio_preset="Personnalisé"
            with cb2:
                if st.button("🔄 Reset",use_container_width=True):
                    if "radio_img" in st.session_state: del st.session_state.radio_img; st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

        # Préréglages
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>⚡ Préréglages</div>", unsafe_allow_html=True)
        presets={"🦴 Os / Fracture":{"mode_radio":True,"clahe":True,"clahe_clip":3.0,"contraste":1.8,"nettete":2.5,"luminosite":1.0,"gamma":0.9},"🦵 Articulaire":{"mode_radio":True,"clahe":True,"clahe_clip":2.0,"contraste":1.5,"nettete":1.8,"luminosite":1.1,"gamma":1.0},"🔲 Négatif radio":{"mode_radio":True,"clahe":True,"clahe_clip":2.0,"contraste":1.5,"nettete":1.5,"negatif":True,"luminosite":1.0,"gamma":1.0},"✨ Haute qualité":{"mode_radio":True,"clahe":True,"clahe_clip":2.5,"contraste":2.0,"nettete":2.0,"debruitage":True,"denoise_r":0.5,"detail_osseux":True,"luminosite":1.0,"gamma":1.0},"☁️ Sous-exposé":{"mode_radio":True,"clahe":True,"clahe_clip":3.0,"contraste":1.6,"luminosite":1.4,"nettete":1.5,"gamma":0.7}}
        for name,pp in presets.items():
            if st.button(name,use_container_width=True):
                if uf:
                    full={"mode_radio":False,"clahe":False,"clahe_clip":2.0,"contraste":1.0,"luminosite":1.0,"nettete":1.0,"gamma":1.0,"debruitage":False,"denoise_r":0.8,"detail_osseux":False,"negatif":False,"egalisation":False}
                    full.update(pp); st.session_state.radio_img=process_radio(Image.open(uf),full); st.session_state.radio_preset=name; st.rerun()
                else: st.warning("Chargez d'abord une image.")
        st.markdown("</div>", unsafe_allow_html=True)

        # Sauvegarde dossier patient
        if uf and opts_p:
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown("<div class='section-title'>💾 Sauvegarder dans dossier patient</div>", unsafe_allow_html=True)
            sel_p_radio=st.selectbox("Patient",list(opts_p.keys()),key="radio_patient_sel")
            region_radio=st.selectbox("Région anatomique",["Rachis lombaire","Rachis cervical","Rachis dorsal","Hanche","Genou","Cheville","Pied","Épaule","Coude","Poignet","Main","Bassin","Fémur","Tibia/Fibula","Humérus","Radius/Cubitus","Thorax","Crâne","Autre"])
            type_radio=st.selectbox("Type de radiographie",["Face","Profil","Face + Profil","3/4","Axiale","Comparative","Autre"])
            notes_radio=st.text_area("Notes cliniques",height=68,placeholder="Observations sur le cliché...")
            if st.button("💾 Sauvegarder dans le dossier",type="primary",use_container_width=True):
                pid_radio=opts_p[sel_p_radio]
                orig_img=Image.open(uf)
                f_orig=save_radio_to_patient(pid_radio,orig_img,region_radio,type_radio,notes_radio,is_traited=False)
                f_traite=None
                if "radio_img" in st.session_state:
                    f_traite=save_radio_to_patient(pid_radio,st.session_state.radio_img,region_radio,type_radio,notes_radio,is_traited=True)
                c.execute("INSERT INTO radios (patient_id,date,region,type_radio,fichier_original,fichier_traite,notes) VALUES (?,?,?,?,?,?,?)",
                          (pid_radio,str(date.today()),region_radio,type_radio,f_orig,f_traite,notes_radio))
                conn.commit(); st.success(f"✅ Radio sauvegardée dans le dossier de {sel_p_radio} !")
            st.markdown("</div>", unsafe_allow_html=True)

    with col_r:
        if uf:
            orig=Image.open(uf)
            c1,c2=st.columns(2)
            with c1:
                st.markdown("<div style='background:#1a3a5c;color:white;padding:0.5rem;border-radius:8px 8px 0 0;text-align:center;font-weight:600;font-size:0.85rem;'>📷 Image originale</div>", unsafe_allow_html=True)
                st.image(orig,use_container_width=True)
                buf_o=io.BytesIO(); orig.save(buf_o,format="PNG")
                st.download_button("⬇️ Télécharger originale",buf_o.getvalue(),"radio_originale.png","image/png",use_container_width=True)
            with c2:
                if "radio_img" in st.session_state:
                    proc=st.session_state.radio_img; label=st.session_state.get("radio_preset","Traitement")
                    st.markdown(f"<div style='background:#065f46;color:white;padding:0.5rem;border-radius:8px 8px 0 0;text-align:center;font-weight:600;font-size:0.85rem;'>✨ {label}</div>", unsafe_allow_html=True)
                    st.image(proc,use_container_width=True)
                    buf_p=io.BytesIO(); proc.save(buf_p,format="PNG")
                    st.download_button("⬇️ Télécharger traitée",buf_p.getvalue(),"radio_traitee.png","image/png",use_container_width=True)
                    # Impression
                    st.markdown("---")
                    st.markdown("**🖨️ Impression :**")
                    buf_print=io.BytesIO(); proc.save(buf_print,format="PNG")
                    b64=__import__('base64').b64encode(buf_print.getvalue()).decode()
                    html_print=f"""<html><head><style>body{{margin:0;padding:0;}} img{{width:100%;height:auto;}} @media print{{body{{margin:0;}}}}</style></head>
<body><img src="data:image/png;base64,{b64}"/><script>window.onload=function(){{window.print();}}</script></body></html>"""
                    b64_html=__import__('base64').b64encode(html_print.encode()).decode()
                    st.markdown(f'<a href="data:text/html;base64,{b64_html}" target="_blank"><button style="width:100%;padding:0.5rem;background:#1a3a5c;color:white;border:none;border-radius:8px;cursor:pointer;font-size:0.9rem;font-weight:600;">🖨️ Imprimer la radio traitée</button></a>', unsafe_allow_html=True)
                else:
                    st.markdown("<div style='background:#f3f4f6;border:2px dashed #d1d5db;border-radius:8px;padding:4rem;text-align:center;color:#9ca3af;'><div style='font-size:3rem;'>🩻</div><p>Cliquez sur Traiter<br>ou choisissez un préréglage</p></div>", unsafe_allow_html=True)
        else:
            st.markdown("<div style='background:#f8fafc;border:2px dashed #d1d5db;border-radius:12px;padding:5rem;text-align:center;color:#9ca3af;margin-top:1rem;'><div style='font-size:4rem;'>🩻</div><h3 style='color:#374151;'>Traitement radiologique</h3><p>Chargez un cliché pour commencer</p></div>", unsafe_allow_html=True)
    conn.close()

# ============================================================
# PAGE : ÉPHARMACIE
# ============================================================
CLASSES_MED=["Tous","AINS","AINS Cox-2","AINS topique","Antalgique","Opioïde faible","Antalgique pallier 2","Myorelaxant","Corticoïde","HBPM","AOD","Bisphosphonate","Supplément","IPP","Antibiotique","Fluoroquinolone","Antioedémateux","Autre"]

def page_pharmacie():
    st.markdown("<div class='main-header'><div style='font-size:2rem;'>💊</div><div><h1>ePharmacie</h1><p>Bibliothèque médicamenteuse complète</p></div></div>", unsafe_allow_html=True)
    tab1,tab2,tab3=st.tabs(["💊 Catalogue","➕ Ajouter","📦 Stocks"])
    conn=get_connection(); c=conn.cursor()
    with tab1:
        cs,cc=st.columns([3,1])
        with cs: search=st.text_input("🔍 Rechercher")
        with cc: cls_f=st.selectbox("Classe",CLASSES_MED)
        q="SELECT * FROM medicaments WHERE 1=1"; p=[]
        if search: q+=" AND (nom LIKE ? OR dci LIKE ? OR classe LIKE ?)"; p+=[f"%{search}%"]*3
        if cls_f!="Tous": q+=" AND classe=?"; p.append(cls_f)
        c.execute(q+" ORDER BY classe,nom",p); meds=c.fetchall()
        st.markdown(f"<p style='color:#6b7280;font-size:0.85rem;'>{len(meds)} médicament(s)</p>", unsafe_allow_html=True)
        classes=list(dict.fromkeys(m["classe"] for m in meds))
        for cls in classes:
            cms=[m for m in meds if m["classe"]==cls]
            st.markdown(f"<div style='background:#1a3a5c;color:white;padding:0.5rem 1rem;border-radius:8px;margin-top:1rem;margin-bottom:0.5rem;font-weight:600;font-size:0.9rem;'>💊 {cls} ({len(cms)})</div>", unsafe_allow_html=True)
            for med in cms:
                sl="❌ Rupture" if med["stock"]==0 else f"⚠️ {med['stock']}" if med["stock"]<10 else f"✅ {med['stock']}"
                with st.expander(f"💊 {med['nom']} — {med['dci']} · {med['dosage']}"):
                    c1,c2=st.columns([3,1])
                    with c1:
                        st.markdown(f"""<table style='width:100%;font-size:0.85rem;'>
                        <tr><td style='color:#6b7280;width:180px;'>DCI</td><td><b>{med['dci']}</b></td></tr>
                        <tr><td style='color:#6b7280;'>Forme/Dosage</td><td>{med['forme']} — {med['dosage']}</td></tr>
                        <tr><td style='color:#6b7280;'>📋 Adulte</td><td style='color:#065f46;font-weight:500;'>{med['posologie_adulte'] or '—'}</td></tr>
                        <tr><td style='color:#6b7280;'>👶 Enfant</td><td>{med['posologie_enfant'] or '—'}</td></tr>
                        <tr><td style='color:#6b7280;'>⛔ Contre-indications</td><td style='color:#dc2626;'>{med['contre_indications'] or '—'}</td></tr>
                        <tr><td style='color:#6b7280;'>⚠️ Effets indésirables</td><td style='color:#d97706;'>{med['effets_indesirables'] or '—'}</td></tr>
                        </table>""", unsafe_allow_html=True)
                    with c2:
                        st.markdown(f"<div style='text-align:center;background:#f8fafc;padding:1rem;border-radius:8px;'><div style='font-size:1.5rem;font-weight:700;color:#1a3a5c;'>{med['prix'] or 0:.0f} DA</div><div style='font-size:0.75rem;color:#6b7280;'>Prix</div><div style='margin-top:0.8rem;font-size:0.82rem;font-weight:600;'>Stock: {sl}</div></div>", unsafe_allow_html=True)
                        if st.button("➕ Ordonnance",key=f"ordo_{med['id']}",use_container_width=True):
                            if "ordo_meds" not in st.session_state: st.session_state.ordo_meds=[]
                            entry={"nom":med["nom"],"dci":med["dci"],"posologie":med["posologie_adulte"],"duree":"À définir","instructions":""}
                            if entry["nom"] not in [x["nom"] for x in st.session_state.ordo_meds]: st.session_state.ordo_meds.append(entry)
                            st.session_state.current_page="ordonnances"; st.rerun()
    with tab2:
        c1,c2=st.columns(2)
        with c1:
            nom=st.text_input("Nom commercial *"); dci=st.text_input("DCI *")
            cls=st.selectbox("Classe *",CLASSES_MED[1:]); forme_m=st.selectbox("Forme",["Comprimé","Comprimé LP","Gélule","Seringue SC","Gel","Sirop","Autre"])
            dos=st.text_input("Dosage")
        with c2:
            prix=st.number_input("Prix (DA)",min_value=0.0,value=0.0,step=10.0)
            stk=st.number_input("Stock",min_value=0,value=0)
            pos_a=st.text_area("Posologie adulte *",height=80); pos_e=st.text_area("Posologie enfant",height=60)
        ci=st.text_area("Contre-indications",height=68); ei=st.text_area("Effets indésirables",height=68)
        if st.button("💾 Enregistrer",type="primary"):
            if nom and dci and pos_a:
                c.execute("INSERT INTO medicaments (nom,dci,classe,forme,dosage,posologie_adulte,posologie_enfant,contre_indications,effets_indesirables,prix,stock) VALUES (?,?,?,?,?,?,?,?,?,?,?)",(nom,dci,cls,forme_m,dos,pos_a,pos_e,ci,ei,prix,stk))
                conn.commit(); st.success(f"✅ {nom} ajouté !"); st.rerun()
            else: st.error("Nom, DCI et posologie adulte obligatoires.")
    with tab3:
        c.execute("SELECT * FROM medicaments ORDER BY stock ASC,nom"); all_m=c.fetchall()
        ca,cb,cc2=st.columns(3)
        with ca: st.metric("Total",len(all_m))
        with cb: st.metric("En rupture",sum(1 for m in all_m if m["stock"]==0))
        with cc2: st.metric("Stock faible",sum(1 for m in all_m if 0<m["stock"]<10))
        st.markdown("---")
        for med in all_m:
            ic="🔴" if med["stock"]==0 else "🟡" if med["stock"]<10 else "🟢"
            c1,c2,c3=st.columns([3,1,1])
            with c1: st.write(f"{ic} **{med['nom']}** ({med['dosage']})")
            with c2: ns=st.number_input("",min_value=0,value=med["stock"],key=f"st_{med['id']}",label_visibility="collapsed")
            with c3:
                if st.button("💾",key=f"sv_{med['id']}"): c.execute("UPDATE medicaments SET stock=? WHERE id=?",(ns,med["id"])); conn.commit(); st.rerun()
    conn.close()

# ============================================================
# PAGE : ORDONNANCES (AMÉLIORÉE — DOCX + IMPRESSION)
# ============================================================
DIAGNOSTICS=["Fracture du radius","Fracture du fémur","Fracture de la cheville","Fracture des côtes",
    "Entorse du genou","Entorse de la cheville","Luxation de l'épaule",
    "Hernie discale lombaire L4-L5","Hernie discale lombaire L5-S1",
    "Cervicalgie / Torticolis","Lombalgie aiguë","Lombalgie chronique",
    "Gonarthrose","Coxarthrose","Ostéoporose","Tendinite rotulienne",
    "Tendinite de la coiffe des rotateurs","Syndrome du canal carpien",
    "Suivi post-PTH","Suivi post-PTG","Contusion musculaire","Déchirure musculaire","Autre"]

PROTOCOLES={"Fracture (douleur post-opératoire)":[
        {"nom":"Paracétamol 1g","dci":"Paracétamol","posologie":"1g toutes les 6h","duree":"7 jours","instructions":""},
        {"nom":"Ibuprofène 400mg","dci":"Ibuprofène","posologie":"400mg 3x/jour pendant les repas","duree":"5 jours","instructions":"Gastroprotection"},
        {"nom":"Oméprazole 20mg","dci":"Oméprazole","posologie":"20mg 1x/jour à jeun","duree":"7 jours","instructions":""}],
    "Lombalgie aiguë":[
        {"nom":"Ibuprofène 400mg","dci":"Ibuprofène","posologie":"400mg 3x/jour avec repas","duree":"5 jours","instructions":""},
        {"nom":"Méthocarbamol 500mg","dci":"Méthocarbamol","posologie":"1500mg 4x/jour","duree":"5 jours","instructions":""},
        {"nom":"Oméprazole 20mg","dci":"Oméprazole","posologie":"20mg 1x/jour à jeun","duree":"5 jours","instructions":""}],
    "Entorse (inflammation aiguë)":[
        {"nom":"Kétoprofène 100mg","dci":"Kétoprofène","posologie":"100mg 2x/jour avec repas","duree":"5 jours","instructions":""},
        {"nom":"Paracétamol 1g","dci":"Paracétamol","posologie":"1g 3x/jour si douleurs","duree":"5 jours","instructions":""},
        {"nom":"Diclofénac gel 1%","dci":"Diclofénac","posologie":"4g 3x/jour local","duree":"7 jours","instructions":"Masser doucement"}],
    "Post-opératoire (anticoagulation)":[
        {"nom":"Énoxaparine 4000UI","dci":"Énoxaparine","posologie":"0.4mL SC 1x/jour","duree":"21 jours","instructions":"Injection sous-cutanée"},
        {"nom":"Paracétamol 1g","dci":"Paracétamol","posologie":"1g 3x/jour","duree":"10 jours","instructions":""},
        {"nom":"Tramadol 50mg","dci":"Tramadol","posologie":"50mg 3x/jour si douleur forte","duree":"5 jours","instructions":""}],
    "Ostéoporose (traitement de fond)":[
        {"nom":"Alendronate 70mg","dci":"Alendronate","posologie":"70mg 1x/semaine à jeun","duree":"3 mois","instructions":"Rester debout 30min après"},
        {"nom":"Calcium + Vitamine D3","dci":"Calcium/Vit D3","posologie":"1 cp 2x/jour après repas","duree":"3 mois","instructions":""}]}

def page_ordonnances():
    st.markdown("<div class='main-header'><div style='font-size:2rem;'>📋</div><div><h1>Ordonnances & Consultations</h1><p>Génération DOCX professionnelle • Impression directe</p></div></div>", unsafe_allow_html=True)
    tab1,tab2,tab3=st.tabs(["📝 Nouvelle consultation","⚡ Ordonnance express","📚 Historique"])
    conn=get_connection(); c=conn.cursor()
    c.execute("SELECT id,nom,prenom,date_naissance,allergies,mutuelle FROM patients ORDER BY nom"); all_p=c.fetchall()
    if not all_p: st.warning("⚠️ Aucun patient enregistré."); conn.close(); return
    opts={f"{p['prenom']} {p['nom']}":dict(p) for p in all_p}

    with tab1:
        c1,c2=st.columns([1,2])
        with c1:
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown("<div class='section-title'>👤 Patient & Consultation</div>", unsafe_allow_html=True)
            def_idx=0
            if "consult_patient_id" in st.session_state:
                for i,(k,v) in enumerate(opts.items()):
                    if v["id"]==st.session_state.consult_patient_id: def_idx=i; break
            sel=st.selectbox("Patient *",list(opts.keys()),index=def_idx)
            pat=opts[sel]
            if pat.get("allergies"):
                st.markdown(f"<div class='alert-warning alert-box'>⚠️ <b>Allergies:</b> {pat['allergies']}</div>", unsafe_allow_html=True)
            dt=st.date_input("Date",value=date.today())
            diag=st.selectbox("Diagnostic",DIAGNOSTICS)
            anam=st.text_area("Anamnèse",height=80)
            exam=st.text_area("Examen clinique",height=80)
            notes_cons=st.text_area("Notes",height=60)
            arret=st.text_input("Arrêt de travail",placeholder="Ex: 7 jours")
            prochain=st.text_input("Prochain RDV",placeholder="Ex: Dans 10 jours")
            inst_gen=st.text_area("Instructions générales",height=68,placeholder="Conseils au patient...")
            st.markdown("</div>", unsafe_allow_html=True)
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown("<div class='section-title'>⚡ Protocoles rapides</div>", unsafe_allow_html=True)
            for pname in PROTOCOLES:
                if st.button(f"📋 {pname}",key=f"proto_{pname}",use_container_width=True):
                    st.session_state.ordo_meds=PROTOCOLES[pname]; st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

        with c2:
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.markdown("<div class='section-title'>💊 Prescription médicamenteuse</div>", unsafe_allow_html=True)
            c.execute("SELECT * FROM medicaments ORDER BY classe,nom"); all_m=c.fetchall()
            med_opts={m["nom"]:dict(m) for m in all_m}
            ms=st.text_input("🔍 Rechercher médicament")
            filt={k:v for k,v in med_opts.items() if not ms or ms.lower() in k.lower() or ms.lower() in v.get("dci","").lower()}
            if filt:
                sel_med=st.selectbox("Médicament",list(filt.keys())); sm=filt[sel_med]
                if sm:
                    st.markdown(f"<div style='background:#f0f9ff;padding:0.6rem;border-radius:6px;font-size:0.82rem;margin-bottom:0.5rem;'><b>Posologie recommandée:</b> {sm['posologie_adulte']}</div>", unsafe_allow_html=True)
                pos=st.text_input("Posologie",value=sm.get("posologie_adulte","") if sm else "")
                dur=st.text_input("Durée",placeholder="Ex: 5 jours, 1 mois")
                inst=st.text_input("Instructions spéciales")
                if st.button("➕ Ajouter à l'ordonnance",type="primary",use_container_width=True):
                    if "ordo_meds" not in st.session_state: st.session_state.ordo_meds=[]
                    st.session_state.ordo_meds.append({"nom":sel_med,"dci":sm.get("dci",""),"posologie":pos,"duree":dur,"instructions":inst}); st.rerun()
            st.markdown("---")
            st.markdown("**📋 Ordonnance en cours:**")
            if "ordo_meds" not in st.session_state: st.session_state.ordo_meds=[]
            if st.session_state.ordo_meds:
                for i,med in enumerate(st.session_state.ordo_meds):
                    cm,cd=st.columns([5,1])
                    with cm: st.markdown(f"<div style='background:#f8fafc;padding:0.6rem;border-radius:6px;margin-bottom:0.3rem;font-size:0.85rem;border-left:3px solid #2d6a9f;'><b>{i+1}. {med['nom']}</b><br>📋 {med['posologie']} · ⏱️ {med['duree']}{('<br>💬 '+med['instructions']) if med.get('instructions') else ''}</div>", unsafe_allow_html=True)
                    with cd:
                        if st.button("🗑️",key=f"del_{i}"): st.session_state.ordo_meds.pop(i); st.rerun()
                if st.button("🗑️ Vider",use_container_width=True): st.session_state.ordo_meds=[]; st.rerun()
                st.markdown("---")

                # Boutons d'action
                ca,cb2,cc2=st.columns(3)
                with ca:
                    if st.button("💾 Sauvegarder",type="primary",use_container_width=True):
                        meds_str="\n".join([f"{m['nom']} — {m['posologie']} — {m['duree']}" for m in st.session_state.ordo_meds])
                        c.execute("INSERT INTO consultations (patient_id,date,anamnese,examen_clinique,diagnostic,traitement,notes) VALUES (?,?,?,?,?,?,?)",(pat["id"],str(dt),anam,exam,diag,meds_str,notes_cons))
                        cid=c.lastrowid
                        c.execute("INSERT INTO ordonnances (patient_id,consultation_id,date,medicaments,duree) VALUES (?,?,?,?,?)",(pat["id"],cid,str(dt),meds_str,"Voir détails"))
                        conn.commit(); st.success("✅ Consultation enregistrée !"); st.session_state.ordo_meds=[]; st.rerun()
                with cb2:
                    if st.button("📄 Générer DOCX",use_container_width=True):
                        with st.spinner("Génération du document Word..."):
                            docx_bytes,err=gen_ordonnance_docx(pat,st.session_state.ordo_meds,dt.strftime("%d/%m/%Y"),diag,inst_gen,prochain,arret)
                        if docx_bytes:
                            st.session_state.ordo_docx=docx_bytes; st.session_state.ordo_docx_name=f"ordonnance_{sel.replace(' ','_')}_{dt}.docx"
                            st.success("✅ Document Word généré !")
                        else: st.error(f"❌ Erreur: {err}")
                with cc2:
                    if "ordo_docx" in st.session_state:
                        st.download_button("⬇️ Télécharger DOCX",st.session_state.ordo_docx,
                            st.session_state.ordo_docx_name,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True)

                # ── Aperçu HTML + Impression (toujours visible) ──
                st.markdown("---")
                entete_h=get_entete()
                rows_html = []
                for i,m in enumerate(st.session_state.ordo_meds):
                    bg = "background:#f8fafc;" if i%2==0 else ""
                    row = (f"<tr style='{bg}'>"
                           f"<td style='padding:8px;border:1px solid #e5e7eb;'><b>{i+1}. {m['nom']}</b>"
                           f"<br><small style='color:#6b7280;'>{m.get('dci','')}</small></td>"
                           f"<td style='padding:8px;border:1px solid #e5e7eb;'>{m['posologie']}</td>"
                           f"<td style='padding:8px;border:1px solid #e5e7eb;'>{m['duree']}</td>"
                           f"<td style='padding:8px;border:1px solid #e5e7eb;color:#92400e;'>{m.get('instructions','')}</td></tr>")
                    rows_html.append(row)
                meds_html = "".join(rows_html)
                allergie_html = ("<div style='color:red;font-weight:bold;'>&#9888; Allergies : " + pat.get('allergies','') + "</div>") if pat.get("allergies") else ""
                rdv_html = ("<b>Prochain RDV :</b> " + prochain + "<br>") if prochain else ""
                arret_html = ("<b style='color:red;'>Arr&ecirc;t de travail :</b> " + arret) if arret else ""
                inst_html = ("<p style='font-style:italic;border-top:1px solid #dbeafe;padding-top:8px;'><b>Instructions :</b> " + inst_gen + "</p>") if inst_gen else "" 
                html_ordo = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><title>Ordonnance</title>
<style>
  body{{font-family:Arial,sans-serif;margin:25px;font-size:13px;color:#111;}}
  .header{{display:flex;justify-content:space-between;align-items:stretch;margin-bottom:10px;}}
  .hdr-left{{flex:1;padding-right:20px;}}
  .hdr-right{{background:#1a3a5c;color:white;padding:12px 18px;border-radius:6px;text-align:center;min-width:200px;}}
  .hdr-right p{{margin:3px 0;font-size:12px;}}
  .sep{{border:none;border-top:3px solid #1a3a5c;margin:12px 0;}}
  .doc-title{{text-align:center;font-size:17px;font-weight:bold;color:#1a3a5c;letter-spacing:2px;margin:14px 0;}}
  .patient-box{{display:flex;justify-content:space-between;background:#f3f4f6;padding:10px 14px;border-radius:6px;margin-bottom:14px;font-size:12px;}}
  table{{width:100%;border-collapse:collapse;margin-top:8px;}}
  thead tr{{background:#1a3a5c;color:white;}}
  th{{padding:8px 10px;text-align:left;font-size:12px;}}
  td{{padding:8px 10px;border:1px solid #e5e7eb;font-size:12px;vertical-align:top;}}
  tr:nth-child(even) td{{background:#f8fafc;}}
  .footer{{display:flex;justify-content:space-between;margin-top:40px;font-size:12px;}}
  .signature{{text-align:center;min-width:220px;}}
  .sig-line{{border-top:1px solid #374151;margin-top:55px;padding-top:5px;color:#9ca3af;font-size:11px;}}
  .pied{{text-align:center;border-top:1px solid #dbeafe;margin-top:20px;padding-top:8px;color:#9ca3af;font-size:10px;}}
  @media print{{
    body{{margin:10px;}}
    .no-print{{display:none;}}
    button{{display:none;}}
  }}
</style></head><body>
<div class="no-print" style="margin-bottom:15px;">
  <button onclick="window.print()" style="padding:10px 28px;background:#1a3a5c;color:white;border:none;border-radius:6px;font-size:14px;cursor:pointer;font-weight:bold;">🖨️ Imprimer</button>
  <button onclick="window.close()" style="padding:10px 20px;background:#6b7280;color:white;border:none;border-radius:6px;font-size:14px;cursor:pointer;margin-left:10px;">✕ Fermer</button>
</div>
<div class="header">
  <div class="hdr-left">
    <div style="font-size:16px;font-weight:bold;color:#1a3a5c;">{entete_h.get("nom_medecin","")}</div>
    <div style="color:#374151;margin-top:3px;">{entete_h.get("specialite","")}</div>
    <div style="color:#6b7280;font-style:italic;font-size:11px;">{entete_h.get("diplomes","")}</div>
  </div>
  <div class="hdr-right">
    <div style="font-weight:bold;font-size:13px;">{entete_h.get("cabinet","")}</div>
    <p>{entete_h.get("adresse","")}</p>
    <p>📞 {entete_h.get("telephone","")}</p>
    <p>{entete_h.get("email","")}</p>
    <p>{entete_h.get("horaires","")}</p>
  </div>
</div>
<hr class="sep">
<div class="doc-title">ORDONNANCE MÉDICALE</div>
<div class="patient-box">
  <div>
    <b style="font-size:13px;">{pat.get("prenom","")} {pat.get("nom","")}</b><br>
    Né(e) le : {pat.get("date_naissance","") or "—"} &nbsp;|&nbsp; Âge : {calc_age(pat.get("date_naissance",""))} ans<br>
    Mutuelle : {pat.get("mutuelle","") or "—"}<br>
    {allergie_html}
  </div>
  <div style="text-align:right;">
    <b>Date :</b> {dt.strftime("%d/%m/%Y")}<br>
    <b>Diagnostic :</b> {diag}
  </div>
</div>
<table>
  <thead><tr><th>Médicament</th><th>Posologie</th><th>Durée</th><th>Instructions</th></tr></thead>
  <tbody>{meds_html}</tbody>
</table>
{inst_html}
<div class="footer">
  <div>{rdv_html}{arret_html}</div>
  <div class="signature">
    <div style="font-weight:bold;color:#1a3a5c;">{entete_h.get("nom_medecin","")}</div>
    <div style="color:#6b7280;font-size:11px;">{entete_h.get("specialite","")}</div>
    <div class="sig-line">Cachet &amp; Signature</div>
  </div>
</div>
<div class="pied">{entete_h.get("cabinet","")} | {entete_h.get("adresse","")} | {entete_h.get("telephone","")} | {entete_h.get("horaires","")}</div>
</body></html>"""
                # Aperçu dans la page
                with st.expander("👁️ Aperçu de l'ordonnance", expanded=False):
                    st.components.v1.html(html_ordo, height=600, scrolling=True)
                # Bouton impression dans nouvel onglet
                b64_html = base64.b64encode(html_ordo.encode("utf-8")).decode()
                st.markdown(
                    f'''<a href="data:text/html;base64,{b64_html}" target="_blank">
                    <button style="width:100%;padding:0.6rem;background:#059669;color:white;
                    border:none;border-radius:8px;cursor:pointer;font-size:0.95rem;
                    font-weight:600;margin-top:6px;">🖨️ Ouvrir et Imprimer l'ordonnance</button></a>''',
                    unsafe_allow_html=True)
            else:
                st.markdown("<div class='alert-info alert-box'>💡 Ajoutez des médicaments ou choisissez un protocole.</div>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

    with tab2:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        sel_q=st.selectbox("Patient",list(opts.keys()),key="qp"); pat_q=opts[sel_q]
        sel_proto=st.selectbox("Protocole thérapeutique",list(PROTOCOLES.keys())); proto_meds=PROTOCOLES[sel_proto]
        for m in proto_meds:
            st.markdown(f"<div style='background:#f0f9ff;padding:0.5rem 1rem;border-radius:6px;margin-bottom:0.3rem;font-size:0.85rem;border-left:3px solid #2563eb;'>💊 <b>{m['nom']}</b> — {m['posologie']} — {m['duree']}</div>", unsafe_allow_html=True)
        col_q1,col_q2=st.columns(2)
        with col_q1:
            if st.button("📄 Générer DOCX",type="primary",use_container_width=True):
                with st.spinner("Génération..."):
                    docx_bytes,err=gen_ordonnance_docx(pat_q,proto_meds,date.today().strftime("%d/%m/%Y"),sel_proto)
                if docx_bytes: st.session_state.quick_docx=docx_bytes; st.session_state.quick_name=f"ordonnance_{sel_q.replace(' ','_')}_{date.today()}.docx"; st.rerun()
                else: st.error(f"❌ {err}")
        with col_q2:
            if "quick_docx" in st.session_state:
                st.download_button("⬇️ Télécharger DOCX",st.session_state.quick_docx,st.session_state.quick_name,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

    with tab3:
        pf=st.selectbox("Filtrer par patient",["Tous"]+list(opts.keys()),key="hfp")
        if pf=="Tous": c.execute("SELECT o.*,p.nom,p.prenom FROM ordonnances o JOIN patients p ON o.patient_id=p.id ORDER BY o.date DESC LIMIT 50")
        else: c.execute("SELECT o.*,p.nom,p.prenom FROM ordonnances o JOIN patients p ON o.patient_id=p.id WHERE o.patient_id=? ORDER BY o.date DESC",(opts[pf]["id"],))
        for o in c.fetchall():
            with st.expander(f"💊 {o['prenom']} {o['nom']} — {o['date']}"):
                for line in (o['medicaments'] or '').split('\n'):
                    if line.strip(): st.write(f"• {line}")
    conn.close()

# ============================================================
# PAGE : RECETTES
# ============================================================
ACTES={"Consultation initiale":2000,"Consultation de suivi":1500,"Consultation d'urgence":2500,"Infiltration articulaire":3000,"Infiltration épidural":4000,"Plâtre / Immobilisation":2500,"Réduction fracture":3500,"Certificat médical":500,"Ponction articulaire":2000,"Pansement / Soin local":800,"Ablation de plâtre":500,"Lecture radiologique":1000,"Consultation pré-opératoire":2000,"Acte autre":0}
MODES=["Espèces","Chèque","CCP","Virement","Mutuelle (tiers payant)","Gratuité"]

def page_recettes():
    st.markdown("<div class='main-header'><div style='font-size:2rem;'>🧾</div><div><h1>Gestion des Recettes</h1><p>Facturation et suivi financier</p></div></div>", unsafe_allow_html=True)
    tab1,tab2,tab3=st.tabs(["➕ Nouvelle recette","📊 Tableau financier","📜 Historique"])
    conn=get_connection(); c=conn.cursor()
    c.execute("SELECT id,nom,prenom FROM patients ORDER BY nom"); all_p=c.fetchall()
    opts={f"{p['prenom']} {p['nom']}":p['id'] for p in all_p}
    with tab1:
        if not all_p: st.warning("Aucun patient."); conn.close(); return
        c1,c2=st.columns(2)
        with c1:
            sel=st.selectbox("👤 Patient *",list(opts.keys()))
            dt=st.date_input("📅 Date",value=date.today())
            acte=st.selectbox("🩺 Acte médical *",list(ACTES.keys()))
            montant=st.number_input("💰 Montant (DA) *",min_value=0.0,value=float(ACTES.get(acte,0)),step=100.0)
        with c2:
            mode=st.selectbox("💳 Mode de paiement",MODES)
            paye=st.checkbox("✅ Paiement reçu",value=True)
            notes=st.text_area("Notes",height=100)
        if st.button("💾 Enregistrer",type="primary",use_container_width=True):
            c.execute("INSERT INTO recettes (patient_id,date,acte,montant,mode_paiement,paye,notes) VALUES (?,?,?,?,?,?,?)",(opts[sel],str(dt),acte,montant,mode,int(paye),notes))
            conn.commit(); st.success(f"✅ Recette de {montant:.0f} DA enregistrée !"); st.rerun()
    with tab2:
        today=date.today()
        cd1,cd2=st.columns(2)
        with cd1: d1=st.date_input("Du",value=today.replace(day=1))
        with cd2: d2=st.date_input("Au",value=today)
        c.execute("SELECT COALESCE(SUM(montant),0) FROM recettes WHERE date=? AND paye=1",(str(today),)); today_t=c.fetchone()[0]
        c.execute("SELECT COALESCE(SUM(montant),0) FROM recettes WHERE strftime('%Y-%m',date)=? AND paye=1",(today.strftime('%Y-%m'),)); month_t=c.fetchone()[0]
        c.execute("SELECT COALESCE(SUM(montant),0) FROM recettes WHERE strftime('%Y',date)=? AND paye=1",(str(today.year),)); year_t=c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM recettes WHERE date BETWEEN ? AND ? AND paye=1",(str(d1),str(d2))); nb=c.fetchone()[0]
        col1,col2,col3,col4=st.columns(4)
        with col1: st.markdown(f"<div class='stat-card'><div class='number'>{today_t:,.0f}</div><div class='label'>DA aujourd'hui</div></div>", unsafe_allow_html=True)
        with col2: st.markdown(f"<div class='stat-card green'><div class='number'>{month_t:,.0f}</div><div class='label'>DA ce mois</div></div>", unsafe_allow_html=True)
        with col3: st.markdown(f"<div class='stat-card orange'><div class='number'>{year_t:,.0f}</div><div class='label'>DA cette année</div></div>", unsafe_allow_html=True)
        with col4: st.markdown(f"<div class='stat-card purple'><div class='number'>{nb}</div><div class='label'>Actes sur période</div></div>", unsafe_allow_html=True)
        c.execute("SELECT r.*,p.nom,p.prenom FROM recettes r JOIN patients p ON r.patient_id=p.id WHERE r.paye=0 ORDER BY r.date DESC"); unpaid=c.fetchall()
        if unpaid:
            total_u=sum(r["montant"] for r in unpaid)
            st.markdown(f"<div class='alert-warning alert-box' style='margin-top:1rem;'>⚠️ <b>{len(unpaid)} facture(s) non payée(s)</b> — Total: <b>{total_u:,.0f} DA</b></div>", unsafe_allow_html=True)
            for r in unpaid:
                cc1,cc2=st.columns([4,1])
                with cc1: st.markdown(f"<div style='background:#fef3c7;padding:0.5rem 1rem;border-radius:6px;font-size:0.85rem;border-left:3px solid #d97706;'>👤 <b>{r['prenom']} {r['nom']}</b> · {r['date']} · {r['acte']} · <b style='color:#d97706;'>{r['montant']:,.0f} DA</b></div>", unsafe_allow_html=True)
                with cc2:
                    if st.button("✅ Payé",key=f"pay_{r['id']}"): c.execute("UPDATE recettes SET paye=1 WHERE id=?",(r["id"],)); conn.commit(); st.rerun()
    with tab3:
        cs2,cp=st.columns([2,2])
        with cs2: sh=st.text_input("🔍 Rechercher")
        with cp: pf=st.selectbox("Patient",["Tous"]+list(opts.keys()),key="hf")
        q="SELECT r.*,p.nom,p.prenom FROM recettes r JOIN patients p ON r.patient_id=p.id WHERE 1=1"; pm=[]
        if pf!="Tous": q+=" AND r.patient_id=?"; pm.append(opts[pf])
        if sh: q+=" AND (p.nom LIKE ? OR p.prenom LIKE ? OR r.acte LIKE ?)"; pm+=[f"%{sh}%"]*3
        c.execute(q+" ORDER BY r.date DESC LIMIT 100",pm)
        for r in c.fetchall():
            pb="<span class='badge badge-green'>✅ Payé</span>" if r['paye'] else "<span class='badge badge-orange'>⏳ En attente</span>"
            st.markdown(f"<div class='card' style='padding:0.8rem 1rem;margin-bottom:0.4rem;'><div style='display:flex;justify-content:space-between;align-items:center;'><div><b>👤 {r['prenom']} {r['nom']}</b> · 📅 {r['date']}<br><span style='font-size:0.82rem;color:#374151;'>🩺 {r['acte']} · 💳 {r['mode_paiement']}</span></div><div style='text-align:right;'><div style='font-size:1.1rem;font-weight:700;color:#1a3a5c;'>{r['montant']:,.0f} DA</div>{pb}</div></div></div>", unsafe_allow_html=True)
    conn.close()

# ============================================================
# PAGE : STATISTIQUES
# ============================================================
def page_statistiques():
    st.markdown("<div class='main-header'><div style='font-size:2rem;'>📊</div><div><h1>Statistiques</h1><p>Analyse de l'activité du cabinet</p></div></div>", unsafe_allow_html=True)
    conn=get_connection(); c=conn.cursor(); today=date.today(); year=today.year; month=today.month
    c.execute("SELECT COUNT(*) FROM patients"); tp=c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM patients WHERE strftime('%Y-%m',created_at)=?",(today.strftime('%Y-%m'),)); npm=c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM consultations"); tc=c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM rendez_vous WHERE statut='terminé'"); tr=c.fetchone()[0]
    col1,col2,col3,col4=st.columns(4)
    with col1: st.metric("Total patients",tp)
    with col2: st.metric("Nouveaux ce mois",npm)
    with col3: st.metric("Consultations totales",tc)
    with col4: st.metric("RDV honorés",tr)
    col_l,col_r=st.columns(2)
    with col_l:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>📅 Activité mensuelle</div>", unsafe_allow_html=True)
        monthly=[]
        for m in range(1,month+1):
            ms=f"{year}-{m:02d}"
            c.execute("SELECT COUNT(*) FROM consultations WHERE strftime('%Y-%m',date)=?",(ms,)); cc=c.fetchone()[0]
            c.execute("SELECT COUNT(*) FROM patients WHERE strftime('%Y-%m',created_at)=?",(ms,)); pc=c.fetchone()[0]
            monthly.append({"mois":calendar.month_abbr[m],"consult":cc,"patients":pc})
        max_c=max((d["consult"] for d in monthly),default=1) or 1
        for d in monthly:
            pct=d["consult"]/max_c*100
            st.markdown(f"<div style='margin-bottom:0.5rem;'><div style='display:flex;justify-content:space-between;font-size:0.82rem;margin-bottom:0.2rem;'><span><b>{d['mois']}</b></span><span>{d['consult']} consultations · {d['patients']} patients</span></div><div style='background:#e5e7eb;border-radius:4px;height:10px;'><div style='background:linear-gradient(90deg,#1a3a5c,#2d6a9f);width:{pct}%;height:10px;border-radius:4px;'></div></div></div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    with col_r:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>💰 Recettes mensuelles</div>", unsafe_allow_html=True)
        rec_data=[]
        for m in range(1,month+1):
            ms=f"{year}-{m:02d}"
            c.execute("SELECT COALESCE(SUM(montant),0) FROM recettes WHERE strftime('%Y-%m',date)=? AND paye=1",(ms,)); t=c.fetchone()[0]
            rec_data.append({"mois":calendar.month_abbr[m],"total":t})
        max_r=max((d["total"] for d in rec_data),default=1) or 1
        for d in rec_data:
            pct=d["total"]/max_r*100
            st.markdown(f"<div style='margin-bottom:0.5rem;'><div style='display:flex;justify-content:space-between;font-size:0.82rem;margin-bottom:0.2rem;'><span><b>{d['mois']}</b></span><span><b>{d['total']:,.0f} DA</b></span></div><div style='background:#e5e7eb;border-radius:4px;height:10px;'><div style='background:linear-gradient(90deg,#065f46,#059669);width:{pct}%;height:10px;border-radius:4px;'></div></div></div>", unsafe_allow_html=True)
        c.execute("SELECT COALESCE(SUM(montant),0) FROM recettes WHERE strftime('%Y',date)=? AND paye=1",(str(year),)); ty=c.fetchone()[0]
        st.markdown(f"<div style='background:#d1fae5;padding:0.8rem;border-radius:8px;margin-top:1rem;text-align:center;'><b>Total {year}: {ty:,.0f} DA</b></div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.markdown("<div class='section-title'>🩺 Diagnostics fréquents</div>", unsafe_allow_html=True)
    c.execute("SELECT diagnostic,COUNT(*) as cnt FROM consultations WHERE diagnostic IS NOT NULL AND diagnostic!='' GROUP BY diagnostic ORDER BY cnt DESC LIMIT 10")
    diags=c.fetchall()
    if diags:
        max_d=diags[0]["cnt"] or 1; colors=["#1a3a5c","#2d6a9f","#3b82f6","#60a5fa","#93c5fd","#bfdbfe"]
        cc1,cc2=st.columns(2)
        for i,d in enumerate(diags):
            pct=d["cnt"]/max_d*100
            with (cc1 if i%2==0 else cc2): st.markdown(f"<div style='margin-bottom:0.6rem;'><div style='display:flex;justify-content:space-between;font-size:0.82rem;margin-bottom:0.2rem;'><span>{d['diagnostic']}</span><span><b>{d['cnt']} cas</b></span></div><div style='background:#e5e7eb;border-radius:4px;height:8px;'><div style='background:{colors[min(i,5)]};width:{pct}%;height:8px;border-radius:4px;'></div></div></div>", unsafe_allow_html=True)
    else: st.info("Aucune donnée.")
    st.markdown("</div>", unsafe_allow_html=True)
    conn.close()

# ============================================================
# PAGE : SAUVEGARDE & RESTAURATION
# ============================================================
def page_sauvegarde():
    st.markdown("<div class='main-header'><div style='font-size:2rem;'>💾</div><div><h1>Sauvegarde & Restauration</h1><p>Protégez toutes vos données patients et radios</p></div></div>", unsafe_allow_html=True)
    col1,col2=st.columns(2)
    with col1:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>📦 Sauvegarder toutes les données</div>", unsafe_allow_html=True)
        st.markdown("""
        <div class='alert-info alert-box'>
            Le fichier ZIP contiendra :<br>
            ✅ Base de données complète (patients, RDV, ordonnances, recettes)<br>
            ✅ Toutes les images radiologiques (originales + traitées)<br>
            ✅ Prêt à être restauré en cas de besoin
        </div>""", unsafe_allow_html=True)
        conn=get_connection(); c=conn.cursor()
        c.execute("SELECT COUNT(*) FROM patients"); np=c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM consultations"); nc=c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM radios"); nr=c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM ordonnances"); no=c.fetchone()[0]
        conn.close()
        st.markdown(f"""
        <div style='display:grid;grid-template-columns:1fr 1fr;gap:0.5rem;margin-bottom:1rem;'>
            <div style='background:#f0f9ff;padding:0.8rem;border-radius:8px;text-align:center;'><b style='font-size:1.3rem;color:#1a3a5c;'>{np}</b><br><small>Patients</small></div>
            <div style='background:#f0f9ff;padding:0.8rem;border-radius:8px;text-align:center;'><b style='font-size:1.3rem;color:#1a3a5c;'>{nc}</b><br><small>Consultations</small></div>
            <div style='background:#f0f9ff;padding:0.8rem;border-radius:8px;text-align:center;'><b style='font-size:1.3rem;color:#1a3a5c;'>{nr}</b><br><small>Radios</small></div>
            <div style='background:#f0f9ff;padding:0.8rem;border-radius:8px;text-align:center;'><b style='font-size:1.3rem;color:#1a3a5c;'>{no}</b><br><small>Ordonnances</small></div>
        </div>""", unsafe_allow_html=True)
        if st.button("📦 Créer la sauvegarde complète", type="primary", use_container_width=True):
            with st.spinner("Création du ZIP en cours..."):
                zip_bytes,zip_name=backup_all_data()
            st.success(f"✅ Sauvegarde créée : {zip_name}")
            st.download_button(f"⬇️ Télécharger {zip_name}", zip_bytes, zip_name, "application/zip", use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

    with col2:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>🗄️ Sauvegardes existantes</div>", unsafe_allow_html=True)
        backups=sorted([f for f in os.listdir(EXPORT_DIR) if f.endswith('.zip')],reverse=True) if os.path.exists(EXPORT_DIR) else []
        if backups:
            for bf in backups[:10]:
                fp=os.path.join(EXPORT_DIR,bf)
                size=os.path.getsize(fp)/1024
                cb1,cb2=st.columns([3,1])
                with cb1: st.write(f"📦 **{bf}** ({size:.0f} KB)")
                with cb2:
                    with open(fp,'rb') as f:
                        st.download_button("⬇️",f.read(),bf,"application/zip",key=f"dl_bk_{bf}")
        else: st.info("Aucune sauvegarde disponible.")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>📤 Export CSV des données</div>", unsafe_allow_html=True)
        conn=get_connection(); c=conn.cursor()
        if st.button("📊 Exporter patients (CSV)"):
            c.execute("SELECT nom,prenom,date_naissance,sexe,telephone,adresse,mutuelle,allergies,antecedents FROM patients")
            rows=c.fetchall()
            csv_lines=["Nom,Prénom,Date Naissance,Sexe,Téléphone,Adresse,Mutuelle,Allergies,Antécédents"]
            for r in rows: csv_lines.append(",".join([f'"{v or ""}"' for v in r]))
            csv_data="\n".join(csv_lines).encode("utf-8-sig")
            st.download_button("⬇️ Télécharger CSV patients",csv_data,f"patients_{date.today()}.csv","text/csv",use_container_width=True)
        if st.button("💰 Exporter recettes (CSV)"):
            c.execute("SELECT p.nom,p.prenom,r.date,r.acte,r.montant,r.mode_paiement,CASE r.paye WHEN 1 THEN 'Payé' ELSE 'En attente' END FROM recettes r JOIN patients p ON r.patient_id=p.id ORDER BY r.date DESC")
            rows=c.fetchall()
            csv_lines=["Nom,Prénom,Date,Acte,Montant,Mode paiement,Statut"]
            for r in rows: csv_lines.append(",".join([f'"{v or ""}"' for v in r]))
            csv_data="\n".join(csv_lines).encode("utf-8-sig")
            st.download_button("⬇️ Télécharger CSV recettes",csv_data,f"recettes_{date.today()}.csv","text/csv",use_container_width=True)
        conn.close()
        st.markdown("</div>", unsafe_allow_html=True)

# ============================================================
# MAIN — ROUTING
# ============================================================
init_db()

if "logged_in" not in st.session_state:
    st.session_state.logged_in=False; st.session_state.role=None; st.session_state.username=None

if not st.session_state.logged_in:
    page_login()
else:
    role=st.session_state.role
    with st.sidebar:
        st.markdown(f"""<div style='background:linear-gradient(135deg,#1a3a5c,#2d6a9f);padding:1.2rem;border-radius:10px;color:white;margin-bottom:1.2rem;'>
            <div style='font-size:1.5rem;'>🦴</div>
            <div style='font-weight:700;font-size:1rem;margin-top:0.3rem;'>Dr. Cabinet</div>
            <div style='font-size:0.75rem;opacity:0.8;'>Traumatologie & Orthopédie</div>
            <div style='background:rgba(255,255,255,0.2);padding:0.2rem 0.6rem;border-radius:20px;font-size:0.7rem;margin-top:0.5rem;display:inline-block;'>{'👨‍⚕️ Médecin' if role=='medecin' else '👩‍💼 Secrétaire'}</div>
        </div>""", unsafe_allow_html=True)

        pages_sec=[("🏠 Accueil","accueil"),("👤 Patients","patients"),("📅 Rendez-vous","rendez_vous"),("🧾 Recettes","recettes"),("💾 Sauvegarde","sauvegarde")]
        pages_med=pages_sec+[("🩻 Radiologie","radiologie"),("💊 ePharmacie","pharmacie"),("📋 Ordonnances","ordonnances"),("📊 Statistiques","statistiques"),("⚙️ En-tête ordonnance","entete")]
        pages=pages_med if role=="medecin" else pages_sec

        if "current_page" not in st.session_state: st.session_state.current_page="accueil"
        st.markdown("**Navigation**")
        for label,key in pages:
            if st.button(label,key=f"nav_{key}",use_container_width=True):
                st.session_state.current_page=key; st.rerun()
        st.markdown("---")
        if st.button("🚪 Déconnexion",use_container_width=True):
            for k in list(st.session_state.keys()): del st.session_state[k]; st.rerun()

    page=st.session_state.current_page
    if   page=="accueil":     page_accueil()
    elif page=="patients":    page_patients()
    elif page=="rendez_vous": page_rendez_vous()
    elif page=="recettes":    page_recettes()
    elif page=="sauvegarde":  page_sauvegarde()
    elif page=="radiologie"   and role=="medecin": page_radiologie()
    elif page=="pharmacie"    and role=="medecin": page_pharmacie()
    elif page=="ordonnances"  and role=="medecin": page_ordonnances()
    elif page=="statistiques" and role=="medecin": page_statistiques()
    elif page=="entete"       and role=="medecin": page_entete()
    else: st.warning("Page non autorisée.")
